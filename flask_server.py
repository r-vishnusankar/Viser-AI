#!/usr/bin/env python3
"""
Flask-based server for Viser AI - more reliable for file uploads
"""
# Fix Windows console encoding so emoji and unicode in print() don't raise UnicodeEncodeError
import sys
import io
if sys.platform == "win32":
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")
    except (AttributeError, OSError):
        pass

from flask import Flask, request, jsonify, send_file, send_from_directory
from flask_cors import CORS
import os
import json
from dotenv import load_dotenv
from pathlib import Path

# Single source of truth: .env only. No config file override for API keys.
_PROJECT_ROOT = Path(__file__).parent
_env_path = _PROJECT_ROOT / ".env"
if _env_path.exists():
    load_dotenv(_env_path)
else:
    load_dotenv()  # fallback to cwd
import uuid
import time
import requests
import traceback
from werkzeug.utils import secure_filename
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from datetime import datetime, date, timedelta
import base64
from io import BytesIO
import threading
import asyncio
from flask_socketio import SocketIO, emit
try:
    import schedule
    SCHEDULE_AVAILABLE = True
except ImportError:
    SCHEDULE_AVAILABLE = False
    print("Warning: schedule library not installed. Calendar scheduler will use simple timer approach.")

# Document generation imports
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Google Sheets integration (live testcase sync)
try:
    from google.oauth2 import service_account
    from googleapiclient.discovery import build
    GOOGLE_SHEETS_AVAILABLE = True
except ImportError:
    GOOGLE_SHEETS_AVAILABLE = False

# Core Engine 2.0 integration
CORE_ENGINE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Core Engine 2.0")
sys.path.insert(0, os.path.join(CORE_ENGINE_DIR, "src"))
sys.path.insert(0, CORE_ENGINE_DIR)  # Allow importing settings.py from Core Engine root

try:
    from spec2.ai_planner import AIPlanner
    from spec2.intent_router import infer_intent
    from spec2.executor import run_with_browser_use
    CORE_ENGINE_AVAILABLE = True
except ImportError:
    CORE_ENGINE_AVAILABLE = False
    print("Warning: spec2 Core Engine modules not found. AI Automation will be disabled.")

# TOON (toonify): compact format to reduce LLM token usage
try:
    from toon import encode as toon_encode
    TOON_AVAILABLE = True
except ImportError:
    TOON_AVAILABLE = False
    toon_encode = None


def tonnify_preprocess(user_prompt: str) -> str:
    """
    TonniFy-style preprocessing: clean and structure user prompt before sending to LLM.
    Reduces junk, normalizes whitespace, keeps only relevant content to lower token usage.
    """
    if not user_prompt or not isinstance(user_prompt, str):
        return user_prompt or ""
    import re
    t = user_prompt.strip()
    t = re.sub(r'\s+', ' ', t)
    t = re.sub(r'^\s*(please|kindly|could you|can you|would you|i want|i need|i would like)\s+', '', t, flags=re.I)
    t = re.sub(r'\s*(please|thanks|thank you|thx)\s*$', '', t, flags=re.I)
    t = t.strip()
    return t if t else user_prompt.strip()

app = Flask(__name__)
app.config['SECRET_KEY'] = 'viser-ai-secret-key'
CORS(app)
socketio = SocketIO(app, cors_allowed_origins="*", async_mode="threading")

# Spec2 Web UI Logger Bridge
class WebUILogger:
    def __init__(self, sio):
        self.sio = sio
    def log(self, level, message):
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.sio.emit('log_message', {
            'timestamp': timestamp,
            'level': level,
            'message': message
        })
        print(f"[{timestamp}] {level}: {message}")

ui_logger = WebUILogger(socketio)
automation_running = False

# Global variables for throttling and context
last_openai_call = 0.0

# Session storage for context awareness
user_sessions = {}
uploaded_files_context = {}

class ContextManager:
    """Manages conversation and file context for users"""
    
    @staticmethod
    def get_session(session_id):
        """Get or create user session, restoring from DB if not in memory."""
        if session_id not in user_sessions:
            history = db_load_session(session_id, limit=60)
            user_sessions[session_id] = {
                "conversation_history": history,
                "uploaded_files": [],
                "last_activity": time.time()
            }
        return user_sessions[session_id]
    
    @staticmethod
    def add_message(session_id, role, content):
        """Add message to conversation history and persist to SQLite."""
        session = ContextManager.get_session(session_id)
        session["conversation_history"].append({
            "role": role,
            "content": content,
            "timestamp": time.time()
        })
        session["last_activity"] = time.time()
        # Persist to DB
        db_save_message(session_id, role, content)

        # Smart context: when history is long, summarise old messages
        if len(session["conversation_history"]) > _SUMMARIZE_THRESHOLD:
            old  = session["conversation_history"][:-_SUMMARIZE_KEEP]
            keep = session["conversation_history"][-_SUMMARIZE_KEEP:]
            summary = _summarize_old_messages(old)
            if summary:
                session["conversation_history"] = [
                    {"role": "system", "content": f"[Earlier conversation summary]\n{summary}",
                     "timestamp": time.time()}
                ] + keep
            else:
                session["conversation_history"] = keep

    @staticmethod
    def add_file_context(session_id, file_info):
        """Add uploaded file to context"""
        session = ContextManager.get_session(session_id)
        session["uploaded_files"].append({
            "filename": file_info["filename"],
            "path": file_info["path"],
            "size": file_info["size"],
            "fileId": file_info.get("fileId", ""),  # Include fileId for frontend
            "type": file_info.get("type", "unknown"),  # Include file type
            "extension": file_info.get("extension", ""),  # Include extension
            "upload_time": time.time(),
            "analyzed": False
        })
        session["last_activity"] = time.time()
    
    @staticmethod
    def get_conversation_context(session_id, include_files=True):
        """Get conversation context for API calls"""
        session = ContextManager.get_session(session_id)
        
        # Build system message with context
        system_message = """You are Viser AI.

TABLE RULES (use tables ONLY when explicitly requested):
- Use a Markdown table ONLY when the user explicitly asks for it (e.g. "in table format", "as a table", "I need it in table", "show as table").
- For all other responses (confirmations, summaries, parsed data, email info, etc.), use normal prose—never tables.
- When the user DOES ask for a table, use this format:
| Column 1 | Column 2 |
| --- | --- |
| Row 1 Data | Row 1 Data |
"""
        
        if include_files and session["uploaded_files"]:
            files_for_context = [
                {"filename": f["filename"], "status": "analyzed" if f["analyzed"] else "uploaded"}
                for f in session["uploaded_files"][-5:]
            ]
            if TOON_AVAILABLE and toon_encode:
                try:
                    context_toon = toon_encode({"uploaded_files": files_for_context})
                    system_message += f"\n\nContext (TOON):\n{context_toon}\nReference these files and offer to analyze if not already done."
                except Exception:
                    file_list = [f"- {f['filename']} ({f['status']})" for f in files_for_context]
                    system_message += f"\n\nContext: User has uploaded files:\n" + "\n".join(file_list) + "\nYou can reference these files in your responses and offer to analyze them if not already done."
            else:
                file_list = [f"- {f['filename']} ({f['status']})" for f in files_for_context]
                system_message += f"\n\nContext: User has uploaded files:\n" + "\n".join(file_list)
                system_message += "\nYou can reference these files in your responses and offer to analyze them if not already done."
        
        # Build message history
        messages = [{"role": "system", "content": system_message}]
        
        # Add recent conversation history (last 10 exchanges)
        recent_history = session["conversation_history"][-20:]  # Last 20 messages (10 exchanges)
        for msg in recent_history:
            messages.append({
                "role": msg["role"],
                "content": msg["content"]
            })
        
        return messages
    
    @staticmethod
    def mark_file_analyzed(session_id, filename):
        """Mark a file as analyzed"""
        session = ContextManager.get_session(session_id)
        for file_info in session["uploaded_files"]:
            if file_info["filename"] == filename:
                file_info["analyzed"] = True
                break
    
    @staticmethod
    def cleanup_old_sessions():
        """Clean up sessions older than 24 hours"""
        current_time = time.time()
        expired_sessions = []
        
        for session_id, session in user_sessions.items():
            if current_time - session["last_activity"] > 86400:  # 24 hours
                expired_sessions.append(session_id)
        
        for session_id in expired_sessions:
            del user_sessions[session_id]

# Automation State
is_automation_running = False

@socketio.on('connect')
def handle_connect():
    ui_logger.log('INFO', 'Client connected to Viser AI Automation Engine')
    emit('status', {'connected': True})

@socketio.on('execute_task')
def handle_execute_task(data):
    """Handle task execution request from web UI"""
    global is_automation_running
    
    if is_automation_running:
        emit('error', {'message': 'An automation task is already running'})
        return
        
    url = data.get('url', '').strip()
    prompt = data.get('prompt', '').strip()
    provider = data.get('provider', 'groq')
    if provider == 'gemini':
        provider = 'groq'
    
    if not url or not prompt:
        emit('error', {'message': 'URL and Objective are required'})
        return

    if not CORE_ENGINE_AVAILABLE:
        emit('error', {'message': 'Core Engine 2.0 not available on this server'})
        return

    # Start task in background
    socketio.start_background_task(run_core_engine_task, url, prompt, provider)

def run_core_engine_task(url, prompt, provider):
    """Execute Core Engine 2.0 task following original logic exactly powder"""
    global is_automation_running
    try:
        is_automation_running = True
        ui_logger.log('INFO', f'🚀 Starting Core Engine 2.0...')
        ui_logger.log('INFO', f'📍 Target URL: {url}')
        ui_logger.log('INFO', f'🎯 Task: {prompt}')
        
        # We need to run the async automation in the background thread
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        
        try:
            from spec2.ai_planner import AIPlanner
            from spec2.intent_router import infer_intent
            from spec2.executor import run_with_browser_use, run_async
            
            # Detect intent
            ui_logger.log('INFO', '🔍 Analyzing intent...')
            intent = infer_intent(prompt)
            ui_logger.log('INFO', f'📋 Detected intent: {intent}')
            
            # Plan the task (fallback to Groq if OpenAI returns 429)
            ui_logger.log('INFO', '🗂️ Planning objectives...')
            planner = AIPlanner(provider)
            plan = planner.plan_for_browser_use(prompt, target_url=url)
            if plan.get('error') and ('429' in str(plan.get('error', '')) or 'insufficient_quota' in str(plan.get('error', '')).lower()) and provider == 'openai' and os.getenv('GROQ_API_KEY'):
                ui_logger.log('INFO', '⚠️ OpenAI quota exceeded, retrying with Groq...')
                planner = AIPlanner('groq')
                plan = planner.plan_for_browser_use(prompt, target_url=url)
            if plan.get('error'):
                ui_logger.log('ERROR', f'❌ Planning failed: {plan["error"]}')
                socketio.emit('error', {'message': f'Planning failed: {plan["error"]}'})
                return
                
            # Emit plan to UI
            socketio.emit('plan_ready', {
                'type': 'single',
                'plan': plan,
                'intent': intent
            })
            
            if plan.get('execution_type') == 'browser_use':
                ui_logger.log('SUCCESS', f'✅ Task planned: {plan.get("task_description", prompt)}')
                ui_logger.log('INFO', '🔄 Executing with AI agent...')
                loop.run_until_complete(run_with_browser_use(url, plan.get('task_description', prompt), socketio, ui_logger))
            else:
                steps = plan.get('steps', [])
                ui_logger.log('SUCCESS', f'✅ Plan created: {len(steps)} steps')
                ui_logger.log('INFO', '🔄 Executing steps with real browser...')
                loop.run_until_complete(run_async(url, steps, socketio, ui_logger))
                
            ui_logger.log('SUCCESS', '✅ Automation task completed successfully')
            socketio.emit('task_success', {'message': 'Task completed successfully'})
            
        except Exception as e:
            ui_logger.log('ERROR', f'💥 Automation failed: {str(e)}')
            socketio.emit('task_error', {'error': str(e)})
            traceback.print_exc()
        finally:
            loop.close()
            
    except Exception as e:
        print(f"❌ Automation crash: {str(e)}")
        socketio.emit('error', {'message': f'Internal engine error: {str(e)}'})
    finally:
        is_automation_running = False
        socketio.emit('task_completed')


@socketio.on('disconnect')
def handle_disconnect():
    ui_logger.log('INFO', '👋 Client disconnected')


@socketio.on('test_connection')
def handle_test_connection():
    emit('log_message', {
        'timestamp': datetime.now().strftime("%H:%M:%S"),
        'level': 'INFO',
        'message': '🔗 Connection OK — Viser AI Core Engine ready',
    })


@socketio.on('plan_task')
def handle_plan_task(data):
    """Generate a plan (dry-run) without executing it."""
    prompt   = (data or {}).get('prompt',   '').strip()
    url      = (data or {}).get('url',      '').strip()
    provider = (data or {}).get('provider', 'groq')
    if provider == 'gemini':
        provider = 'groq'

    if not prompt:
        emit('error', {'message': 'Please provide a task prompt'})
        return

    if not CORE_ENGINE_AVAILABLE:
        emit('error', {'message': 'Core Engine 2.0 not available on this server'})
        return

    socketio.start_background_task(_plan_task_bg, prompt, url, provider)


def _plan_task_bg(prompt: str, url: str, provider: str) -> None:
    """Background task: run AI planning without browser execution."""
    try:
        from spec2.ai_planner import AIPlanner, compare
        from spec2.intent_router import infer_intent

        ui_logger.log('INFO', f'🎯 Planning: {prompt}')
        if url:
            ui_logger.log('INFO', f'🌐 Target URL: {url}')
        ui_logger.log('INFO', f'🤖 Provider: {provider}')

        intent = infer_intent(prompt)
        ui_logger.log('INFO', f'📋 Intent: {intent}')

        if provider == 'compare':
            _plan_compare_bg(prompt, url, intent)
            return

        planner = AIPlanner(provider)
        plan = planner.plan(prompt, url)
        if plan.get('error') and ('429' in str(plan.get('error', '')) or 'insufficient_quota' in str(plan.get('error', '')).lower()) and provider == 'openai' and os.getenv('GROQ_API_KEY'):
            ui_logger.log('INFO', '⚠️ OpenAI quota exceeded, retrying with Groq...')
            planner = AIPlanner('groq')
            plan = planner.plan(prompt, url)
        if plan.get('error'):
            ui_logger.log('ERROR', f'❌ Planning failed: {plan["error"]}')
            socketio.emit('error', {'message': f'Planning failed: {plan["error"]}'})
            return

        steps_count = len(plan.get('steps', []))
        ui_logger.log('SUCCESS', f'✅ Plan ready: {steps_count} steps')

        if plan.get('saved_to'):
            ui_logger.log('INFO', f'📁 Saved: {plan["saved_to"]}')

        socketio.emit('plan_ready', {'type': 'single', 'plan': plan, 'intent': intent})

    except Exception as exc:
        ui_logger.log('ERROR', f'💥 Planning error: {exc}')
        socketio.emit('error', {'message': f'Planning failed: {exc}'})


def _plan_compare_bg(prompt: str, url: str, intent: str) -> None:
    """Compare plans from all available providers side-by-side."""
    try:
        from spec2.ai_planner import compare
        ui_logger.log('INFO', '🔄 Comparing all available providers...')
        results = compare(prompt, url)

        for prov, result in results.items():
            if result.get('error'):
                ui_logger.log('ERROR', f'❌ {prov.upper()}: {result["error"]}')
            else:
                n = len(result.get('steps', []))
                ui_logger.log('INFO', f'✅ {prov.upper()}: {n} steps')

        socketio.emit('plan_ready', {'type': 'compare', 'results': results, 'intent': intent})

    except Exception as exc:
        ui_logger.log('ERROR', f'💥 Compare failed: {exc}')
        socketio.emit('error', {'message': f'Compare failed: {exc}'})


@socketio.on('enhance_plan')
def handle_enhance_plan(data):
    """Enhance an uploaded plan file via AI."""
    raw      = (data or {}).get('raw',      '').strip()
    filename = (data or {}).get('filename', 'uploaded_plan.txt')
    provider = (data or {}).get('provider', 'groq')
    if provider == 'gemini':
        provider = 'groq'
    url      = (data or {}).get('url',      '').strip()

    if not raw:
        emit('error', {'message': 'No plan content received'})
        return

    if not CORE_ENGINE_AVAILABLE:
        emit('error', {'message': 'Core Engine 2.0 not available on this server'})
        return

    socketio.start_background_task(_enhance_plan_bg, raw, filename, provider, url)


def _enhance_plan_bg(raw: str, filename: str, provider: str, url: str) -> None:
    """Background task: use AI to improve an uploaded plan."""
    try:
        from spec2.ai_planner import AIPlanner
        ui_logger.log('INFO', f'📂 Enhancing plan: {filename}')
        planner = AIPlanner(provider)
        prompt = (
            "Please read and improve this execution plan. Convert it into specific, "
            "actionable browser automation steps. Fix structure, rephrase steps, and "
            f"return detailed JSON with a steps list.\n\nFile: {filename}\n\nContent:\n\n{raw}"
        )
        plan = planner.plan(prompt, url)
        if plan.get('error') and ('429' in str(plan.get('error', '')) or 'insufficient_quota' in str(plan.get('error', '')).lower()) and provider == 'openai' and os.getenv('GROQ_API_KEY'):
            ui_logger.log('INFO', '⚠️ OpenAI quota exceeded, retrying with Groq...')
            planner = AIPlanner('groq')
            plan = planner.plan(prompt, url)
        if plan.get('error'):
            ui_logger.log('ERROR', f'❌ Enhancement failed: {plan["error"]}')
            socketio.emit('error', {'message': f'Enhancement failed: {plan["error"]}'})
            return

        if plan.get('saved_to'):
            ui_logger.log('INFO', f'📁 Enhanced plan saved: {plan["saved_to"]}')

        socketio.emit('plan_ready', {'type': 'single', 'plan': plan, 'intent': 'uploaded'})
        ui_logger.log('SUCCESS', f'✅ Plan enhanced: {filename}')

    except Exception as exc:
        ui_logger.log('ERROR', f'💥 Enhancement error: {exc}')
        socketio.emit('error', {'message': f'Failed to enhance plan: {exc}'})


@app.route('/api/repo/files', methods=['GET'])
def list_repo_files():
    """List all available repository files in the Test Archive folder"""
    try:
        repo_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Test Archive')
        if not os.path.exists(repo_dir):
            return jsonify({"files": []})
            
        files = [f for f in os.listdir(repo_dir) if f.endswith('.xlsx')]
        return jsonify({
            "success": True,
            "files": files,
            "count": len(files)
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/repo/load/<filename>', methods=['GET'])
def load_repo_file(filename):
    """Return the raw file for client-side parsing"""
    try:
        repo_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Test Archive')
        return send_from_directory(repo_dir, filename)
    except Exception as e:
        return jsonify({"error": str(e)}), 404


def _extract_sheet_id(url_or_id):
    """Extract spreadsheet ID from URL or return as-is if already an ID."""
    s = (url_or_id or '').strip()
    if not s:
        return None
    # URL format: https://docs.google.com/spreadsheets/d/SPREADSHEET_ID/edit...
    if 'spreadsheets/d/' in s:
        start = s.find('spreadsheets/d/') + len('spreadsheets/d/')
        end = len(s)
        for i in range(start, len(s)):
            if s[i] in '/?':
                end = i
                break
        return s[start:end].strip()
    # Assume it's already an ID (alphanumeric, dashes)
    return s


@app.route('/api/repo/gsheets', methods=['GET'])
def repo_gsheets():
    """Fetch testcase data from Google Sheets. Requires GOOGLE_SHEETS_CREDENTIALS_JSON env var."""
    if not GOOGLE_SHEETS_AVAILABLE:
        return jsonify({"success": False, "error": "Google Sheets API not available. Install google-auth and google-api-python-client."}), 503
    sheet_id = request.args.get('sheet_id') or request.args.get('url')
    if not sheet_id:
        return jsonify({"success": False, "error": "Missing sheet_id or url parameter"}), 400
    spreadsheet_id = _extract_sheet_id(sheet_id)
    if not spreadsheet_id:
        return jsonify({"success": False, "error": "Invalid sheet URL or ID"}), 400
    creds_json = os.environ.get('GOOGLE_SHEETS_CREDENTIALS_JSON')
    if not creds_json:
        return jsonify({"success": False, "error": "GOOGLE_SHEETS_CREDENTIALS_JSON not configured"}), 503
    try:
        creds_dict = json.loads(creds_json)
    except json.JSONDecodeError:
        try:
            creds_dict = json.loads(base64.b64decode(creds_json).decode('utf-8'))
        except Exception:
            return jsonify({"success": False, "error": "Invalid GOOGLE_SHEETS_CREDENTIALS_JSON"}), 500
    try:
        credentials = service_account.Credentials.from_service_account_info(creds_dict)
        service = build('sheets', 'v4', credentials=credentials)
        spreadsheet = service.spreadsheets().get(spreadsheetId=spreadsheet_id).execute()
        sheet_names = [s['properties']['title'] for s in spreadsheet.get('sheets', [])]
        all_rows = []
        for sheet_name in sheet_names:
            result = service.spreadsheets().values().get(
                spreadsheetId=spreadsheet_id,
                range=f"'{sheet_name}'!A:Z"
            ).execute()
            values = result.get('values', [])
            if not values:
                continue
            headers = values[0]
            for row in values[1:]:
                obj = {}
                for i, h in enumerate(headers):
                    obj[h] = row[i] if i < len(row) else ''
                all_rows.append(obj)
        return jsonify({
            "success": True,
            "data": all_rows,
            "source": "gsheets",
            "sheets": sheet_names
        })
    except Exception as e:
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/unsplash/random', methods=['GET'])
def unsplash_random():
    """Fetch a random Unsplash image URL for Nexora background. Requires UNSPLASH_ACCESS_KEY in .env"""
    key = os.environ.get('UNSPLASH_ACCESS_KEY')
    if not key:
        return jsonify({"enabled": False})
    try:
        query = request.args.get('query', 'gradient abstract purple')
        r = requests.get(
            'https://api.unsplash.com/photos/random',
            params={'query': query, 'orientation': 'landscape', 'client_id': key},
            timeout=5
        )
        r.raise_for_status()
        data = r.json()
        return jsonify({
            "enabled": True,
            "url": data.get("urls", {}).get("regular"),
            "thumb": data.get("urls", {}).get("thumb"),
        })
    except Exception as e:
        return jsonify({"enabled": True, "error": str(e)}), 500


def extract_docx_content(file_path):
    """Extract text content from .docx files"""
    try:
        from docx import Document
        doc = Document(file_path)
        content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text.strip())
        return '\n'.join(content)
    except ImportError:
        return "Error: python-docx library not installed. Cannot read .docx files."
    except Exception as e:
        return f"Error reading .docx file: {str(e)}"

def extract_pdf_content(file_path):
    """Extract text content from PDF files"""
    try:
        import PyPDF2
        content = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                content.append(page.extract_text())
        return '\n'.join(content)
    except ImportError:
        return "Error: PyPDF2 library not installed. Cannot read PDF files."
    except Exception as e:
        return f"Error reading PDF file: {str(e)}"

def extract_image_content(file_path):
    """Extract base64 content from image files for AI analysis"""
    try:
        import base64
        
        # Read image file and convert to base64
        with open(file_path, 'rb') as image_file:
            image_data = image_file.read()
            base64_data = base64.b64encode(image_data).decode('utf-8')
            
            # Determine MIME type based on file extension
            file_ext = os.path.splitext(file_path)[1].lower()
            mime_types = {
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg', 
                '.png': 'image/png',
                '.gif': 'image/gif',
                '.bmp': 'image/bmp',
                '.webp': 'image/webp',
                '.tiff': 'image/tiff',
                '.svg': 'image/svg+xml'
            }
            
            mime_type = mime_types.get(file_ext, 'image/jpeg')
            
            return {
                "base64_data": base64_data,
                "mime_type": mime_type,
                "file_size": len(image_data)
            }
            
    except Exception as e:
        print(f"❌ Image extraction error: {str(e)}")
        return None

def analyze_image_with_vision(file_path, prompt="Analyze this image and describe what you see"):
    """Analyze image using OpenAI Vision API (Gemini removed)"""
    try:
        import base64
        image_content = extract_image_content(file_path)
        if not image_content:
            return "❌ Failed to extract image content"
        openai_key = CONFIG.get("OPENAI_API_KEY")
        if not openai_key:
            return "❌ OpenAI API key not configured for image analysis"
        import openai as _oai
        client = _oai.OpenAI(api_key=openai_key)
        with open(file_path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode()
        mime = image_content.get("mime_type", "image/jpeg")
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}}
                ]
            }],
            max_tokens=1024
        )
        content = response.choices[0].message.content
        print(f"✅ Image analysis completed: {len(content)} characters")
        return content
    except Exception as e:
        error_msg = f"Image analysis error: {str(e)}"
        print(f"❌ {error_msg}")
        return f"❌ {error_msg}"

def get_fallback_analysis(files):
    """Generate fallback analysis when AI is unavailable"""
    if isinstance(files, list):
        names = ", ".join([f.get('name', 'file') for f in files])
    else:
        names = files if isinstance(files, str) else "uploaded file"
    
    return (
        f"**Fallback Analysis for: {names}**\n\n"
        "⚠️ AI analysis temporarily unavailable. Here's a basic assessment:\n\n"
        "**Identified Issues:**\n"
        "- 3 pricing conflicts found across platforms\n"
        "- 2 inventory discrepancies detected\n"
        "- 5 product description inconsistencies\n\n"
        "**Key Metrics to Monitor:**\n"
        "- Average Order Value (AOV)\n"
        "- Click-Through Rate (CTR)\n"
        "- Sell-through Rate\n\n"
        "**Recommended Actions:**\n"
        "- Standardize product descriptions\n"
        "- Adjust pricing for consistency\n"
        "- Restock fast-moving inventory\n"
        "- Review and update product categories"
    )

def convert_messages_for_gemini(messages):
    """Convert OpenAI/Groq message format to Gemini format"""
    gemini_messages = []
    system_message = ""
    
    for msg in messages:
        if msg["role"] == "system":
            system_message = msg["content"]
        elif msg["role"] == "user":
            content = msg["content"]
            if system_message:
                content = f"System: {system_message}\n\nUser: {content}"
                system_message = ""  # Only add system message once
            gemini_messages.append({
                "role": "user",
                "parts": [{"text": content}]
            })
        elif msg["role"] == "assistant":
            gemini_messages.append({
                "role": "model",
                "parts": [{"text": msg["content"]}]
            })
    
    return gemini_messages

def detect_email_command(message):
    """Detect if message is an email send command and extract details"""
    import re
    
    # Pattern 1: "send [message] to [email]" - improved to handle longer messages
    pattern1 = r'send\s+(.+?)\s+to\s+([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})'
    match1 = re.search(pattern1, message, re.IGNORECASE | re.DOTALL)
    
    if match1:
        message_content = match1.group(1).strip()
        recipient_email = match1.group(2).strip()
        
        # If message content is just "this" or "it", send previous AI response
        if message_content.lower() in ['this', 'it']:
            return {
                'is_email_command': True,
                'message_content': 'AI Response',  # Will be replaced with actual content
                'recipient_email': recipient_email,
                'send_previous_response': True
            }
        else:
            return {
                'is_email_command': True,
                'message_content': message_content,
                'recipient_email': recipient_email
            }
    
    # Pattern 2: Just an email address (for sending previous AI response)
    pattern2 = r'^([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})$'
    match2 = re.search(pattern2, message.strip())
    
    if match2:
        return {
            'is_email_command': True,
            'message_content': 'AI Response',  # Will be replaced with actual content
            'recipient_email': match2.group(1).strip(),
            'send_previous_response': True
        }
    
    return {'is_email_command': False}

def create_simple_email_body(message_content, session_id=None):
    """Create plain email body for quick messages (no template)"""
    return f"<p>{message_content}</p>"

def send_email(to_email, subject, body, attachment_path=None, attachment_name=None):
    """Send email with optional attachment"""
    if not CONFIG.get('EMAIL_ENABLED', False):
        print("📧 Email sending disabled in configuration")
        return False, "Email sending is disabled. Set EMAIL_ENABLED=True in .env"

    # Validate required config
    sender = CONFIG.get('SENDER_EMAIL', '').strip()
    app_pwd = CONFIG.get('APP_PASSWORD', '').strip()
    smtp_server = CONFIG.get('SMTP_SERVER', '').strip()
    smtp_port = CONFIG.get('SMTP_PORT', 587)

    if not sender:
        return False, "SENDER_EMAIL is not configured. Add your email to .env"
    if not app_pwd:
        return False, "APP_PASSWORD is not configured. For Gmail, use an App Password (not your regular password). Generate at: https://myaccount.google.com/apppasswords"
    if not smtp_server:
        return False, "SMTP_SERVER is not configured in .env"
    if not to_email or not str(to_email).strip():
        return False, "Recipient email address is required"

    to_email = str(to_email).strip()

    try:
        # Create message
        msg = MIMEMultipart()
        msg['From'] = sender
        msg['To'] = to_email
        msg['Subject'] = subject

        # Add body
        msg.attach(MIMEText(body, 'html'))

        # Add attachment if provided
        if attachment_path and os.path.exists(attachment_path):
            with open(attachment_path, "rb") as attachment:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(attachment.read())

            encoders.encode_base64(part)
            part.add_header(
                'Content-Disposition',
                f'attachment; filename= {attachment_name or os.path.basename(attachment_path)}',
            )
            msg.attach(part)

        # Send email (use sendmail for reliability across providers)
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.ehlo()
        server.starttls()
        server.ehlo()
        server.login(sender, app_pwd)
        server.sendmail(sender, to_email, msg.as_string())
        server.quit()

        print(f"✅ Email sent successfully to {to_email}")
        return True, "Email sent successfully"

    except smtplib.SMTPAuthenticationError as e:
        err = "SMTP authentication failed. For Gmail: use an App Password (not your regular password). Enable 2FA and generate at: https://myaccount.google.com/apppasswords"
        print(f"❌ {err}: {e}")
        return False, err
    except smtplib.SMTPException as e:
        err = f"SMTP error: {str(e)}"
        print(f"❌ {err}")
        return False, err
    except Exception as e:
        error_msg = f"Failed to send email: {str(e)}"
        print(f"❌ {error_msg}")
        return False, error_msg

def create_analysis_email_body(filename, analysis, session_info=None):
    """Create plain email body for analysis report (no template)"""
    parts = [f"<p><strong>File:</strong> {filename}</p>", f"<pre>{analysis}</pre>"]
    if session_info:
        parts.append(f"<p>{session_info}</p>")
    return "\n".join(parts)

def create_notification_email_body(title, message, details=None):
    """Create plain email body for notifications (no template)"""
    parts = [f"<p><strong>{title}</strong></p>", f"<p>{message}</p>"]
    if details:
        parts.append(f"<p>{details}</p>")
    return "\n".join(parts)

def parse_markdown_table(content):
    """Parse markdown table content into structured data"""
    lines = content.strip().split('\n')
    table_lines = []
    
    for line in lines:
        line = line.strip()
        if line.startswith('|'):
            # Tolerant of missing trailing pipe (Gemini often omits it)
            rest = line[1:].rstrip('|').strip()
            cells = [cell.strip() for cell in rest.split('|')]
            table_lines.append(cells)
    
    if len(table_lines) < 2:
        return None, None
    
    # First line is headers, rest are data
    headers = table_lines[0]
    
    # Filter out separator lines (rows that only contain dashes/colons)
    data_rows = []
    for row in table_lines[1:]:
        is_separator = all(all(c in ' -:|' for c in cell) for cell in row)
        if not is_separator:
            data_rows.append(row)
    
    return headers, data_rows

def generate_pdf_document(content, filename="document.pdf"):
    """Generate PDF document from content"""
    if not REPORTLAB_AVAILABLE:
        return None, "PDF generation requires reportlab library"
    
    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=1*inch)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        story.append(Paragraph("Vise-AI Generated Document", title_style))
        story.append(Spacer(1, 20))
        
        # Check if content contains tables
        headers, data_rows = parse_markdown_table(content)
        
        if headers and data_rows:
            # Create table
            table_data = [headers] + data_rows
            
            # Create table with styling
            table = Table(table_data, repeatRows=1)
            table.setStyle(TableStyle([
                # Header styling
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                
                # Data styling
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            
            story.append(table)
        else:
            # Regular text content
            for line in content.split('\n'):
                if line.strip():
                    story.append(Paragraph(line.strip(), styles['Normal']))
                    story.append(Spacer(1, 6))
        
        # Footer
        story.append(Spacer(1, 20))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.grey,
            alignment=1  # Center alignment
        )
        story.append(Paragraph(f"Generated by Vise-AI on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", footer_style))
        
        doc.build(story)
        buffer.seek(0)
        
        return buffer.getvalue(), None
        
    except Exception as e:
        return None, f"PDF generation error: {str(e)}"

def generate_docx_document(content, filename="document.docx"):
    """Generate DOCX document from content"""
    if not DOCX_AVAILABLE:
        return None, "DOCX generation requires python-docx library"
    
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading('Vise-AI Generated Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add timestamp
        timestamp = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Check if content contains tables
        headers, data_rows = parse_markdown_table(content)
        
        if headers and data_rows:
            # Create table
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            
            # Add headers
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                # Style header cells
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add data rows
            for row_data in data_rows:
                row_cells = table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    row_cells[i].text = cell_data
        else:
            # Regular text content
            for line in content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue(), None
        
    except Exception as e:
        return None, f"DOCX generation error: {str(e)}"

def generate_excel_document(content, filename="document.xlsx"):
    """Generate Excel document from content"""
    if not OPENPYXL_AVAILABLE:
        return None, "Excel generation requires openpyxl library"
    
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Vise-AI Document"
        
        # Add title
        ws['A1'] = "Vise-AI Generated Document"
        ws['A2'] = f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        
        # Style title
        ws['A1'].font = Font(size=16, bold=True)
        ws['A2'].font = Font(size=10, italic=True)
        
        # Check if content contains tables
        headers, data_rows = parse_markdown_table(content)
        
        if headers and data_rows:
            # Start table from row 4
            start_row = 4
            
            # Add headers
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=start_row, column=col, value=header)
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill(start_color="667eea", end_color="667eea", fill_type="solid")
                cell.alignment = Alignment(horizontal="center", vertical="center")
            
            # Add data rows
            for row_idx, row_data in enumerate(data_rows, start_row + 1):
                for col_idx, cell_data in enumerate(row_data, 1):
                    cell = ws.cell(row=row_idx, column=col_idx, value=cell_data)
                    cell.alignment = Alignment(wrap_text=True, vertical="top")
                    
                    # Alternate row colors
                    if row_idx % 2 == 0:
                        cell.fill = PatternFill(start_color="f8fafc", end_color="f8fafc", fill_type="solid")
            
            # Auto-adjust column widths
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width
        else:
            # Regular text content
            for row_idx, line in enumerate(content.split('\n'), 4):
                if line.strip():
                    ws.cell(row=row_idx, column=1, value=line.strip())
        
        # Save to buffer
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue(), None
        
    except Exception as e:
        return None, f"Excel generation error: {str(e)}"

def generate_txt_document(content, filename="document.txt"):
    """Generate plain text document from content"""
    try:
        # Clean up content for plain text
        lines = content.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line and not line.startswith('|---'):  # Skip markdown table separators
                cleaned_lines.append(line)
        
        # Add header and footer
        document_content = []
        document_content.append("=" * 60)
        document_content.append("VISE-AI GENERATED DOCUMENT")
        document_content.append("=" * 60)
        document_content.append(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        document_content.append("")
        document_content.extend(cleaned_lines)
        document_content.append("")
        document_content.append("=" * 60)
        document_content.append("End of Document")
        document_content.append("=" * 60)
        
        return '\n'.join(document_content), None
        
    except Exception as e:
        return None, f"Text generation error: {str(e)}"

# Basic config - loads from environment variables
CONFIG = {
    "owner_name": os.getenv("OWNER_NAME", "Vishnu"),
    # Which LLM answers chat and analysis: "groq" | "openai" | "gemini" | "fallback"
    "AI_PROVIDER": os.getenv("AI_PROVIDER", "groq"),
    "OPENAI_API_KEY": os.getenv("OPENAI_API_KEY", ""),
    "OPENAI_MODEL": os.getenv("OPENAI_MODEL", "o4-mini-2025-04-16"),
    "GROQ_API_KEY": os.getenv("GROQ_API_KEY", ""),
    "GROQ_MODEL": os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile"),
    "GEMINI_API_KEY": os.getenv("GEMINI_API_KEY", ""),
    "GEMINI_MODEL": os.getenv("GEMINI_MODEL", "gemini-2.0-flash"),
    # Seconds to wait for LLM API response (increase if you get read timeout)
    "API_TIMEOUT": int(os.getenv("API_TIMEOUT", "120")),

    # Email Configuration
    "EMAIL_ENABLED": os.getenv("EMAIL_ENABLED", "True").lower() == "true",
    "SMTP_SERVER": os.getenv("SMTP_SERVER", "smtp.gmail.com"),
    "SMTP_PORT": int(os.getenv("SMTP_PORT", "587")),
    "SENDER_EMAIL": os.getenv("SENDER_EMAIL", ""),
    "APP_PASSWORD": os.getenv("APP_PASSWORD", ""),
    "OWNER_EMAIL": os.getenv("OWNER_EMAIL", ""),
    "DEFAULT_RECIPIENT": os.getenv("DEFAULT_RECIPIENT", ""),
}

def get_ai_provider():
    """
    Return current LLM provider: groq or openai only. Gemini removed.
    Uses Groq by default; falls back to OpenAI if Groq key missing.
    """
    p = os.environ.get("AI_PROVIDER", CONFIG.get("AI_PROVIDER", "groq")).strip().lower()
    if p == "gemini":
        p = "groq"
    if p not in ("groq", "openai"):
        p = "groq" if CONFIG.get("GROQ_API_KEY") else "openai"
    return p


# ─── SQLite Chat Persistence ───────────────────────────────────────────────────
import sqlite3

_DB_PATH = os.path.join(os.path.dirname(__file__), "data", "chat_history.db")

def _get_db():
    os.makedirs(os.path.dirname(_DB_PATH), exist_ok=True)
    conn = sqlite3.connect(_DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("""
        CREATE TABLE IF NOT EXISTS chat_messages (
            id        INTEGER PRIMARY KEY AUTOINCREMENT,
            session_id TEXT NOT NULL,
            role      TEXT NOT NULL,
            content   TEXT NOT NULL,
            ts        REAL NOT NULL
        )
    """)
    conn.execute("CREATE INDEX IF NOT EXISTS idx_session ON chat_messages(session_id, ts)")
    conn.commit()
    return conn

def db_save_message(session_id: str, role: str, content: str):
    try:
        conn = _get_db()
        conn.execute("INSERT INTO chat_messages(session_id,role,content,ts) VALUES(?,?,?,?)",
                     (session_id, role, content, time.time()))
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"⚠️ DB save error: {e}")

def db_load_session(session_id: str, limit: int = 60):
    """Load the most recent `limit` messages for a session from DB."""
    try:
        conn = _get_db()
        rows = conn.execute(
            "SELECT role,content,ts FROM chat_messages WHERE session_id=? ORDER BY ts DESC LIMIT ?",
            (session_id, limit)
        ).fetchall()
        conn.close()
        return [{"role": r["role"], "content": r["content"], "timestamp": r["ts"]}
                for r in reversed(rows)]
    except Exception as e:
        print(f"⚠️ DB load error: {e}")
        return []

def db_clear_session(session_id: str):
    """Clear all messages for a session from DB."""
    try:
        conn = _get_db()
        conn.execute("DELETE FROM chat_messages WHERE session_id=?", (session_id,))
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"⚠️ DB clear error: {e}")

def db_clear_session(session_id: str):
    try:
        conn = _get_db()
        conn.execute("DELETE FROM chat_messages WHERE session_id=?", (session_id,))
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"⚠️ DB clear error: {e}")

def db_all_sessions():
    """Return list of {session_id, count, last_ts} for the history view."""
    try:
        conn = _get_db()
        rows = conn.execute("""
            SELECT session_id, COUNT(*) as cnt, MAX(ts) as last_ts
            FROM chat_messages GROUP BY session_id ORDER BY last_ts DESC LIMIT 50
        """).fetchall()
        conn.close()
        return [{"session_id": r["session_id"], "count": r["cnt"], "last_ts": r["last_ts"]} for r in rows]
    except Exception as e:
        print(f"⚠️ DB sessions error: {e}")
        return []

# Bootstrap DB on import
try:
    _get_db().close()
except Exception:
    pass

# ─── Smart Context Summarisation ──────────────────────────────────────────────
_SUMMARIZE_THRESHOLD = 30   # summarize when history exceeds this many messages
_SUMMARIZE_KEEP      = 10   # keep the N most recent messages verbatim

def _summarize_old_messages(old_messages: list) -> str:
    """Call the active AI to produce a short summary of older conversation turns."""
    if not old_messages:
        return ""
    try:
        combined = "\n".join(
            f"{m['role'].upper()}: {m['content'][:300]}" for m in old_messages
        )
        prompt = (
            "Summarize this conversation so far in 3-5 bullet points. "
            "Be concise; preserve important facts, decisions, and files mentioned.\n\n"
            + combined
        )
        provider = get_ai_provider()
        if provider == "openai":
            import openai as _oai
            _client = _oai.OpenAI(api_key=CONFIG["OPENAI_API_KEY"])
            r = _client.chat.completions.create(
                model=CONFIG.get("OPENAI_MODEL", "o4-mini-2025-04-16"),
                messages=[{"role":"user","content":prompt}],
                max_tokens=300, temperature=0.3
            )
            return r.choices[0].message.content.strip()
        else:  # groq (default)
            r = requests.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers={"Authorization": f"Bearer {CONFIG['GROQ_API_KEY']}", "Content-Type":"application/json"},
                json={"model": CONFIG.get("GROQ_MODEL","llama-3.3-70b-versatile"),
                      "messages":[{"role":"user","content":prompt}],
                      "max_tokens":300,"temperature":0.3},
                timeout=30
            )
            return r.json()["choices"][0]["message"]["content"].strip()
    except Exception as e:
        print(f"⚠️ Smart context summarise error: {e}")
        return ""

# Calendar Events Data Storage
CALENDAR_EVENTS_FILE = "data/calendar_events.json"
CALENDAR_IMAGES_DIR = "uploads/calendar"

def ensure_calendar_directories():
    """Ensure calendar data directories exist"""
    os.makedirs("data", exist_ok=True)
    os.makedirs(CALENDAR_IMAGES_DIR, exist_ok=True)

def load_calendar_events():
    """Load calendar events from JSON file"""
    ensure_calendar_directories()
    if os.path.exists(CALENDAR_EVENTS_FILE):
        try:
            with open(CALENDAR_EVENTS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f"Error loading calendar events: {e}")
            return []
    return []

def save_calendar_events(events):
    """Save calendar events to JSON file"""
    ensure_calendar_directories()
    try:
        with open(CALENDAR_EVENTS_FILE, 'w', encoding='utf-8') as f:
            json.dump(events, f, indent=2, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Error saving calendar events: {e}")
        return False

def create_calendar_event(event_date, recipient_email, image_path, event_type="anniversary", message=""):
    """Create a new calendar event"""
    events = load_calendar_events()
    event_id = str(uuid.uuid4())
    
    event = {
        "id": event_id,
        "date": event_date,  # Format: YYYY-MM-DD
        "email": recipient_email,
        "image_path": image_path,
        "event_type": event_type,
        "message": message,
        "sent": False,
        "created_at": datetime.now().isoformat(),
        "sent_at": None
    }
    
    events.append(event)
    save_calendar_events(events)
    return event

def get_events_for_date(target_date):
    """Get all events for a specific date"""
    events = load_calendar_events()
    target_date_str = target_date.strftime("%Y-%m-%d") if isinstance(target_date, date) else target_date
    
    return [e for e in events if e["date"] == target_date_str and not e.get("sent", False)]

def mark_event_as_sent(event_id):
    """Mark an event as sent"""
    events = load_calendar_events()
    for event in events:
        if event["id"] == event_id:
            event["sent"] = True
            event["sent_at"] = datetime.now().isoformat()
            save_calendar_events(events)
            return True
    return False

def create_calendar_email_body(message, event_type, image_path=None):
    """Create HTML email body for calendar events"""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    image_html = ""
    if image_path and os.path.exists(image_path):
        # Embed image in email
        with open(image_path, 'rb') as img_file:
            img_data = base64.b64encode(img_file.read()).decode('utf-8')
            img_ext = os.path.splitext(image_path)[1].lower()
            mime_type = {
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.png': 'image/png',
                '.gif': 'image/gif'
            }.get(img_ext, 'image/jpeg')
            
            image_html = f'<div style="text-align: center; margin: 20px 0;"><img src="data:{mime_type};base64,{img_data}" alt="Calendar Event Image" style="max-width: 100%; height: auto; border-radius: 8px; box-shadow: 0 4px 8px rgba(0,0,0,0.1);"></div>'
    
    html_body = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; line-height: 1.6; color: #333; }}
            .header {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 20px; text-align: center; }}
            .content {{ padding: 30px; background: #f8f9fa; }}
            .message {{ background: white; padding: 20px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); font-size: 16px; }}
            .footer {{ background: #343a40; color: white; padding: 15px; text-align: center; font-size: 12px; }}
            .event-type {{ background: #667eea; color: white; padding: 5px 15px; border-radius: 20px; display: inline-block; margin-bottom: 15px; font-size: 14px; font-weight: 500; }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>🎉 {event_type.title()} Reminder</h1>
        </div>
        <div class="content">
            <div class="event-type">{event_type.title()}</div>
            {image_html}
            <div class="message">
                {message if message else f'<p>Wishing you a wonderful {event_type}!</p>'}
            </div>
        </div>
        <div class="footer">
            <p>Sent via Vise-AI on {timestamp}</p>
        </div>
    </body>
    </html>
    """
    return html_body

def send_calendar_event_email(event):
    """Send email for a calendar event with image"""
    try:
        # Create email body
        body = create_calendar_email_body(
            event.get("message", ""),
            event.get("event_type", "anniversary"),
            event.get("image_path")
        )
        
        # Create subject
        subject = f"🎉 {event.get('event_type', 'Anniversary').title()} Reminder"
        
        # Send email
        success, message = send_email(
            event["email"],
            subject,
            body,
            attachment_path=event.get("image_path") if event.get("image_path") else None
        )
        
        if success:
            mark_event_as_sent(event["id"])
            print(f"✅ Calendar event email sent to {event['email']} for {event['date']}")
            return True
        else:
            print(f"❌ Failed to send calendar event email: {message}")
            return False
    except Exception as e:
        print(f"❌ Error sending calendar event email: {str(e)}")
        return False

def check_and_send_calendar_events():
    """Check for events today and send emails"""
    today = date.today()
    events = get_events_for_date(today)
    
    if events:
        print(f"📅 Found {len(events)} calendar events for today ({today})")
        for event in events:
            send_calendar_event_email(event)
    else:
        print(f"📅 No calendar events for today ({today})")

def run_scheduler():
    """Run the scheduler in a background thread"""
    def scheduler_thread():
        if SCHEDULE_AVAILABLE:
            # Use schedule library for precise timing
            schedule.every().day.at("09:00").do(check_and_send_calendar_events)
            
            print("⏰ Calendar scheduler started - checking daily at 9:00 AM")
            
            while True:
                schedule.run_pending()
                time.sleep(60)  # Check every minute
        else:
            # Simple timer approach - check every hour, send at 9 AM
            print("⏰ Calendar scheduler started (simple mode) - checking daily at 9:00 AM")
            
            while True:
                now = datetime.now()
                current_hour = now.hour
                current_minute = now.minute
                
                # Check if it's 9 AM (within 1 minute window)
                if current_hour == 9 and current_minute == 0:
                    check_and_send_calendar_events()
                    # Wait 1 minute to avoid sending multiple times
                    time.sleep(60)
                else:
                    # Check every minute
                    time.sleep(60)
    
    thread = threading.Thread(target=scheduler_thread, daemon=True)
    thread.start()
    return thread

# Core Engine 2.0 Automation Dashboard
@app.route('/automation')
@app.route('/automation/')
def serve_automation():
    """Serve the Core Engine 2.0 automation dashboard (web_ui_v2 template)."""
    template_path = os.path.join(CORE_ENGINE_DIR, 'templates_webui_v2', 'index.html')
    if os.path.exists(template_path):
        response = send_file(template_path)
        response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
        return response
    return "Core Engine dashboard template not found. Ensure 'Core Engine 2.0/templates_webui_v2/index.html' exists.", 404


# Enhanced static file serving
@app.route('/')
@app.route('/index.html')
def serve_index():
    try:
        response = send_file('viser-ai-modern.html')
        # Prevent caching so UI updates are always visible after refresh
        response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
        return response
    except FileNotFoundError:
        return "HTML file not found", 404

@app.route('/assets/<path:subpath>')
def serve_assets(subpath):
    """Serve assets (images, etc.) from assets/ folder"""
    try:
        assets_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'assets')
        return send_from_directory(assets_dir, subpath)
    except Exception as e:
        print(f"❌ Asset serve error: {e}")
        return "Not found", 404

@app.route('/<path:filename>')
def serve_static(filename):
    """Serve static files (HTML, CSS, JS)"""
    try:
        # Security: only allow certain file types
        allowed_extensions = {'.html', '.css', '.js', '.png', '.jpg', '.jpeg', '.gif', '.ico'}
        file_ext = os.path.splitext(filename)[1].lower()
        
        if file_ext not in allowed_extensions:
            return "File type not allowed", 403
            
        # Set appropriate content type
        content_types = {
            '.html': 'text/html',
            '.css': 'text/css', 
            '.js': 'application/javascript',
            '.png': 'image/png',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.gif': 'image/gif',
            '.ico': 'image/x-icon'
        }
        
        if os.path.exists(filename):
            response = send_file(filename)
            response.headers['Content-Type'] = content_types.get(file_ext, 'text/plain')
            return response
        else:
            return "File not found", 404
            
    except Exception as e:
        print(f"❌ Static file error: {e}")
        return "Error serving file", 500

@app.route('/api/chat', methods=['POST'])
def chat():
    try:
        print("📨 Received chat request")
        data = request.get_json()
        user_message = data.get("message", "")
        session_id = data.get("session_id", "default_session")
        
        print(f"💬 User message: {user_message[:100]}...")
        print(f"🔑 Session ID: {session_id}")

        if not user_message:
            print("❌ No message provided")
            return jsonify({"error": "No message provided"}), 400

        # Add user message to conversation history (TonniFy preprocessed for token efficiency)
        cleaned = tonnify_preprocess(user_message)
        ContextManager.add_message(session_id, "user", cleaned if cleaned else user_message)
        
        # Clean up old sessions periodically
        if len(user_sessions) > 100:
            ContextManager.cleanup_old_sessions()

        # Check if this is an email send command
        email_command = detect_email_command(user_message)
        if email_command['is_email_command']:
            # Handle sending previous AI response if just email address provided
            if email_command.get('send_previous_response'):
                # Get the last AI response from conversation history
                session = ContextManager.get_session(session_id)
                last_ai_response = None
                
                # Find the last assistant message
                for msg in reversed(session["conversation_history"]):
                    if msg["role"] == "assistant":
                        last_ai_response = msg["content"]
                        break
                
                if last_ai_response:
                    message_content = last_ai_response
                    subject = f"AI Response from Vise-AI"
                else:
                    message_content = "No previous AI response found to send."
                    subject = f"Message from Vise-AI"
            else:
                message_content = email_command['message_content']
                subject = f"Message from Vise-AI User"
            
            # Send the email
            body = create_simple_email_body(
                message_content,
                session_id
            )
            
            success, message = send_email(
                email_command['recipient_email'],
                subject,
                body
            )
            
            if success:
                ai_response = f"✅ Email sent successfully to {email_command['recipient_email']}!\n\nMessage: \"{message_content[:100]}{'...' if len(message_content) > 100 else ''}\""
            else:
                ai_response = f"❌ Failed to send email: {message}"
            
            ContextManager.add_message(session_id, "assistant", ai_response)
            return jsonify({
                "response": ai_response,
                "session_id": session_id,
                "email_sent": success
            })

        # Check if using fallback mode
        if get_ai_provider() == 'fallback':
            ai_response = "Vise-AI fallback: Please provide your data files or a specific BI question."
            return jsonify({"response": ai_response})

        # OpenAI throttling (if using OpenAI)
        global last_openai_call
        if get_ai_provider() == 'openai':
            gap = 5.0 - (time.time() - last_openai_call)
            if gap > 0:
                print(f"⏳ OpenAI throttling: waiting {gap:.2f} seconds")
                time.sleep(gap)

        # Get conversation context with history
        messages = ContextManager.get_conversation_context(session_id, include_files=True)
        print(f"📚 Context: {len(messages)} messages in conversation")

        # --- API call for chat (OpenAI, Groq, or Gemini based on config) ---
        if get_ai_provider() == 'openai':
            headers = {
                "Authorization": f"Bearer {CONFIG['OPENAI_API_KEY']}",
                "Content-Type": "application/json"
            }
            api_url = "https://api.openai.com/v1/chat/completions"
            model = CONFIG.get("OPENAI_MODEL", "o4-mini-2025-04-16")
            
            payload = {
                "model": model,
                "messages": messages,
                "temperature": 0.6,
                "max_tokens": 8192
            }
            
            print(f"🔄 Making OpenAI API call to {api_url}...")
            response = requests.post(api_url, headers=headers, json=payload, timeout=CONFIG.get('API_TIMEOUT', 120))
            print(f"📊 API Response Status: {response.status_code}")
            
            if response.status_code == 200:
                result = response.json()
                ai_response = result["choices"][0]["message"]["content"]
                last_openai_call = time.time()  # Update throttling timestamp
            else:
                raise Exception(f"OpenAI API Error {response.status_code}: {response.text}")
                
        else:  # Default to Groq
            headers = {
                "Authorization": f"Bearer {CONFIG['GROQ_API_KEY']}",
                "Content-Type": "application/json"
            }
            api_url = "https://api.groq.com/openai/v1/chat/completions"
            model = CONFIG.get("GROQ_MODEL", "llama-3.3-70b-versatile")
            
            payload = {
                "model": model,
                "messages": messages,
                "temperature": 0.6,
                "max_tokens": 8192
            }
            
            print(f"🔄 Making Groq API call to {api_url}...")
            response = requests.post(api_url, headers=headers, json=payload, timeout=CONFIG.get('API_TIMEOUT', 120))
            print(f"📊 Groq API Response Status: {response.status_code}")
            
            if response.status_code == 200:
                result = response.json()
                ai_response = result["choices"][0]["message"]["content"]
            else:
                raise Exception(f"Groq API Error {response.status_code}: {response.text}")

        print(f"✅ AI Response generated: {len(ai_response)} characters")
        
        # --- Table enforcement only when user explicitly asks for a table ---
        asks_for_table = any(p in user_message.lower() for p in (
            "in table", "as a table", "as table", "in tabular", "tabular format",
            "table format", "show as table", "need it in table", "format as table"
        ))
        if asks_for_table:
            headers_table, rows_table = parse_markdown_table(ai_response)
            if headers_table is None or not rows_table:
                # Auto-retry instead of scolding the user
                print("⚠️ Invalid table detected, retrying...")
                retry_prompt = (
                    "Please provide ONLY a valid Markdown table "
                    "with headers, one separator row, and real data rows."
                )
                messages = ContextManager.get_conversation_context(session_id)
                messages.append({"role": "user", "content": retry_prompt})

                # Retry with the same active provider (not hard-coded Gemini)
                active_provider = get_ai_provider()
                retry_response = None

                if active_provider == 'openai':
                    retry_response = requests.post(
                        "https://api.openai.com/v1/chat/completions",
                        headers={"Authorization": f"Bearer {CONFIG['OPENAI_API_KEY']}", "Content-Type": "application/json"},
                        json={"model": CONFIG.get("OPENAI_MODEL", "o4-mini-2025-04-16"), "messages": messages, "temperature": 0.3, "max_tokens": 4096},
                        timeout=CONFIG.get('API_TIMEOUT', 120)
                    )
                    if retry_response.status_code == 200:
                        ai_response = retry_response.json()["choices"][0]["message"]["content"]

                else:  # groq (default)
                    retry_response = requests.post(
                        "https://api.groq.com/openai/v1/chat/completions",
                        headers={"Authorization": f"Bearer {CONFIG['GROQ_API_KEY']}", "Content-Type": "application/json"},
                        json={"model": CONFIG.get("GROQ_MODEL", "llama-3.3-70b-versatile"), "messages": messages, "temperature": 0.3, "max_tokens": 4096},
                        timeout=CONFIG.get('API_TIMEOUT', 120)
                    )
                    if retry_response.status_code == 200:
                        ai_response = retry_response.json()["choices"][0]["message"]["content"]

                if retry_response and retry_response.status_code == 200:
                    
                    # --- ADVANCED STRUCTURAL REPAIR ---
                    def fix_table_hallucination(text):
                        if "|" not in text or "---" not in text: return text
                        
                        # AI likes to use | --- | as a row separator on one line
                        # Convert Header | --- | Data1 | --- | Data2 -> Proper Table
                        import re
                        
                        # 1. Unroll multiple pipes
                        text = text.replace("||", "|\n|")
                        
                        # 2. Fix cases where header, separator, and data are condensed
                        # "Header | --- | Data" -> "Header\n| --- |\nData"
                        text = re.sub(r'(\| [^|\n]+ \|)\s*(\|[\s:-]*\|)\s*(\| [^|\n]+ \|)', r'\1\n\2\n\3', text)
                        
                        # 3. Fix cases where AI uses | --- | as a newline between data rows
                        # "| D1 | D2 | | --- | | D3 | D4 |" or "| D1 | D2 | --- | D3 | D4 |"
                        text = re.sub(r'(\| [^|\n]+ \|)\s*(\|[\s:-]*\|)\s*(\| [^|\n]+ \|)', r'\1\n\2\n\3', text)
                        
                        # 4. Final Cleanup: Ensure every line starting with | is separated by a newline
                        lines = []
                        for line in text.strip().split('\n'):
                            if "|" in line and "---" not in line and line.count("|") > 4:
                                # This looks like a header + data combo
                                if "---" in line:
                                    parts = line.split("| --- |")
                                    lines.extend([p.strip() + "|" if not p.strip().endswith("|") else p.strip() for p in parts])
                                else:
                                    lines.append(line)
                            else:
                                lines.append(line)
                        return "\n".join(lines)

                    ai_response = fix_table_hallucination(ai_response)
                    print("✅ Advanced Table structural repair completed")
                    
                    print("✅ Table retry successful")
        
        # Add AI response to conversation history
        ContextManager.add_message(session_id, "assistant", ai_response)
        
        return jsonify({
            "response": ai_response,
            "session_id": session_id
        })

    except requests.exceptions.RequestException as e:
        error_msg = f"API request failed: {str(e)}"
        print(f"❌ Request error: {error_msg}")
        return jsonify({"error": error_msg, "type": "network_error"}), 500
    except Exception as e:
        print(f"❌ Unexpected error: {str(e)}")
        traceback.print_exc()
        return jsonify({"error": f"Unexpected error: {str(e)}"}), 500

def get_file_type(extension):
    """Determine file type based on extension"""
    ext = extension.lower()
    
    # Image types
    image_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.tiff', '.svg'}
    
    # Document types
    doc_extensions = {'.pdf', '.doc', '.docx', '.txt', '.rtf', '.odt'}
    
    if ext in image_extensions:
        return 'image'
    elif ext in doc_extensions:
        return 'document'
    else:
        return 'other'

@app.route('/api/upload', methods=['POST'])
def upload():
    try:
        print("📤 Received upload request")
        
        # Get session ID from form data or default
        session_id = request.form.get('session_id', 'default_session')
        print(f"🔑 Session ID: {session_id}")
        
        if 'file' not in request.files:
            print("❌ No file in request")
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        if file.filename == '':
            print("❌ No file selected")
            return jsonify({"error": "No file selected"}), 400
        
        print(f"📄 File: {file.filename} ({file.content_type})")
        
        # Read file data
        file_data = file.read()
        print(f"📊 Read {len(file_data)} bytes")
        
        # Determine file type
        file_extension = os.path.splitext(file.filename)[1] if '.' in file.filename else '.txt'
        file_type = get_file_type(file_extension)
        print(f"📂 File type: {file_type}")
        
        # Create appropriate uploads directory
        if file_type == 'image':
            upload_dir = "uploads/images"
        else:
            upload_dir = "uploads/documents"
        os.makedirs(upload_dir, exist_ok=True)
        
        # Generate unique filename
        file_id = str(uuid.uuid4())
        unique_filename = f"{file_id}{file_extension}"
        file_path = os.path.join(upload_dir, unique_filename)
        
        # Save file
        with open(file_path, 'wb') as f:
            f.write(file_data)
        
        print(f"✅ Upload successful: {file.filename} -> {unique_filename}")
        
        # Add file to context
        file_info = {
            "filename": file.filename,
            "path": file_path,
            "size": len(file_data),
            "fileId": file_id,
            "type": file_type,
            "extension": file_extension.lower()
        }
        ContextManager.add_file_context(session_id, file_info)
        print(f"📚 Added file to context for session: {session_id}")
        
        return jsonify({
            "success": True,
            "message": f"File uploaded successfully: {file.filename}",
            "fileId": file_id,
            "filename": file.filename,
            "size": len(file_data),
            "path": file_path,
            "type": file_type,
            "extension": file_extension.lower(),
            "session_id": session_id
        })
        
    except Exception as e:
        print(f"❌ Upload error: {str(e)}")
        traceback.print_exc()
        return jsonify({"error": f"Upload failed: {str(e)}"}), 500

@app.route('/api/analyze', methods=['POST'])
def analyze():
    try:
        print("🔍 Received analysis request")
        data = request.get_json()
        filename = data.get("filename")
        file_path = data.get("file_path")
        session_id = data.get("session_id", "default_session")
        
        print(f"📄 Analyzing: {filename}")
        print(f"📂 File path: {file_path}")
        print(f"🔑 Session ID: {session_id}")

        # Determine file type and read content accordingly
        if file_path and os.path.exists(file_path):
            try:
                file_ext = os.path.splitext(filename)[1].lower()
                file_type = get_file_type(file_ext)
                
                if file_type == 'image':
                    print("🖼️ Analyzing image with Gemini Vision...")
                    # For images, analyze directly with Gemini Vision API
                    analysis = analyze_image_with_vision(file_path, "Analyze this image in detail and describe what you see")
                    
                    # Mark file as analyzed in context
                    ContextManager.mark_file_analyzed(session_id, filename)
                    ContextManager.add_message(session_id, "assistant", f"Image Analysis: {analysis}")
                    
                    return jsonify({
                        "filename": filename,
                        "analysis": analysis,
                        "file_type": "image",
                        "session_id": session_id
                    })
                    
                elif filename.lower().endswith('.docx'):
                    print("📄 Extracting DOCX content...")
                    content = extract_docx_content(file_path)
                elif filename.lower().endswith('.pdf'):
                    print("📄 Extracting PDF content...")
                    content = extract_pdf_content(file_path)
                elif filename.lower().endswith(('.txt', '.md', '.csv')):
                    print("📄 Reading text file...")
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                else:
                    print("📄 Reading file as text...")
                    # For other file types, try to read as text
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
            except Exception as e:
                content = f"Error reading file content: {str(e)}"
                print(f"❌ File reading error: {e}")
        else:
            return jsonify({"error": "File not found or no file path provided"}), 400

        # Check if using fallback mode
        if get_ai_provider() == 'fallback':
            print("🔄 Using fallback analysis mode")
            fallback_response = get_fallback_analysis(filename)
            return jsonify({
                "filename": filename,
                "analysis": fallback_response
            })
        
        # Increase content length limit for more detailed analysis
        # Use a smarter approach: keep more content but warn if very large
        max_content_length = 50000  # Increased from 8000 to allow more detailed analysis
        original_length = len(content)
        content_truncated = False
        
        if len(content) > max_content_length:
            # Try to keep important parts - take from beginning and end
            half_limit = max_content_length // 2
            content = content[:half_limit] + "\n\n[... CONTENT TRUNCATED ...]\n\n" + content[-half_limit:]
            content_truncated = True
            print(f"📏 Content truncated from {original_length} to ~{max_content_length} characters (keeping start and end)")
        else:
            print(f"📏 Analyzing full document content: {original_length} characters")
        
        # Determine document type for specialized analysis
        doc_type_hint = ""
        if any(keyword in filename.lower() for keyword in ['test', 'case', 'scenario', 'spec']):
            doc_type_hint = "This appears to be a test case or specification document."
        elif any(keyword in filename.lower() for keyword in ['api', 'endpoint', 'rest', 'swagger']):
            doc_type_hint = "This appears to be an API documentation or specification."
        elif any(keyword in filename.lower() for keyword in ['readme', 'guide', 'manual', 'tutorial']):
            doc_type_hint = "This appears to be a guide or documentation."
        elif any(keyword in filename.lower() for keyword in ['code', 'script', 'program', '.py', '.js', '.java']):
            doc_type_hint = "This appears to be a code file or technical document."
        
        # Build enhanced analysis prompt for detailed analysis
        prompt = f"""
You are Vise-AI, an expert document analyst with deep expertise in technical documentation, business documents, and various content types. 
Your task is to provide a COMPREHENSIVE, DETAILED, and THOROUGH analysis of the following document.

{doc_type_hint}

**Document Filename:** {filename}
**Content Length:** {original_length} characters{' (truncated for analysis)' if content_truncated else ' (full content)'}

**Document Content:**
{content}

---

## REQUIRED ANALYSIS SECTIONS (Provide detailed information for each):

### 1. DOCUMENT OVERVIEW & METADATA
- **Primary Purpose:** What is the main objective of this document?
- **Target Audience:** Who is this document intended for? (developers, managers, end-users, etc.)
- **Document Type:** Classification (technical spec, user guide, test case, API doc, etc.)
- **Main Topics:** List all major topics and themes covered
- **Document Structure:** How is the document organized? (sections, chapters, etc.)

### 2. DETAILED CONTENT ANALYSIS
- **Key Concepts:** Explain all major concepts, ideas, and principles discussed
- **Important Information:** List all critical facts, figures, dates, and data points
- **Methodologies:** Describe any processes, workflows, or methodologies outlined
- **Frameworks & Standards:** Identify any frameworks, standards, or best practices mentioned
- **Terminology:** Define important terms, acronyms, and technical jargon used

### 3. TECHNICAL DETAILS & SPECIFICATIONS
- **Requirements:** List all requirements, prerequisites, and dependencies
- **Configuration:** Document all configuration options, settings, and parameters
- **Code Examples:** Extract and explain any code snippets, examples, or scripts
- **Technical Specifications:** Detail any technical specifications, limits, or constraints
- **Integration Points:** Identify APIs, endpoints, interfaces, or integration details

### 4. PRACTICAL IMPLEMENTATION INFORMATION
- **Step-by-Step Processes:** Provide detailed walkthroughs of any procedures
- **Instructions:** List all actionable instructions and how-to guidance
- **Best Practices:** Extract and explain recommended practices
- **Common Pitfalls:** Identify warnings, cautions, or common mistakes mentioned
- **Troubleshooting:** Document any troubleshooting steps or solutions provided

### 5. DATA & METRICS
- **Key Metrics:** Extract any numbers, statistics, or measurements
- **Tables & Lists:** Summarize information from tables, lists, or structured data
- **Examples:** Detail all examples, use cases, or scenarios provided
- **References:** List any external references, links, or related documents

### 6. BUSINESS & STRATEGIC INSIGHTS
- **Business Context:** Explain business implications and context
- **Use Cases:** Describe practical applications and use cases
- **Benefits & Value:** Identify benefits, advantages, or value propositions
- **Risks & Considerations:** Note any risks, limitations, or important considerations

### 7. COMPREHENSIVE SUMMARY
- **Executive Summary:** Provide a high-level summary (2-3 paragraphs)
- **Key Takeaways:** List the 5-10 most important points from the document
- **Action Items:** Extract any actionable items, next steps, or recommendations
- **Important Notes:** Highlight any critical information that must be remembered
- **Related Topics:** Suggest related topics or areas for further exploration

---

## ANALYSIS GUIDELINES:
- Be THOROUGH and DETAILED - don't skip important information
- Provide SPECIFIC examples and quotes from the document where relevant
- Use clear formatting with bullet points, numbered lists, and sections
- Use COMPACT formatting: NO extra blank lines between sections or bullet points. Use at most ONE blank line between major sections. Keep all spacing minimal.
- Include ALL relevant technical details, not just summaries
- Explain complex concepts in detail
- Extract and preserve important data points, numbers, and specifications
- If the document contains code, explain what it does in detail
- If the document describes processes, provide step-by-step breakdowns
- Be comprehensive - aim for a detailed analysis that captures the full value of the document

Begin your detailed analysis now:
"""

        # --- API call for analysis (OpenAI, Groq, or Gemini based on config) ---
        analysis_messages = [
            {"role": "system", "content": "You are Vise-AI, an expert document analyst. Provide detailed, comprehensive analysis. Use COMPACT formatting: minimal blank lines, tight spacing between sections and bullet points. No extra vertical space."},
            {"role": "user", "content": prompt}
        ]
        
        if get_ai_provider() == 'openai':
            headers = {
                "Authorization": f"Bearer {CONFIG['OPENAI_API_KEY']}",
                "Content-Type": "application/json"
            }
            api_url = "https://api.openai.com/v1/chat/completions"
            model = CONFIG.get("OPENAI_MODEL", "o4-mini-2025-04-16")
            
            payload = {
                "model": model,
                "messages": analysis_messages,
                "temperature": 0.3,
                "max_tokens": 4000  # Increased from 2000 to allow more detailed responses
            }
            
            print(f"🔄 Making OpenAI analysis API call...")
            response = requests.post(api_url, headers=headers, json=payload, timeout=CONFIG.get('API_TIMEOUT', 120))
            print(f"📊 Analysis API Response Status: {response.status_code}")
            
            if response.status_code == 200:
                result = response.json()
                ai_response = result["choices"][0]["message"]["content"]
            else:
                raise Exception(f"OpenAI Analysis API Error {response.status_code}: {response.text}")
                
        else:  # Default to Groq
            headers = {
                "Authorization": f"Bearer {CONFIG['GROQ_API_KEY']}",
                "Content-Type": "application/json"
            }
            api_url = "https://api.groq.com/openai/v1/chat/completions"
            model = CONFIG.get("GROQ_MODEL", "llama-3.3-70b-versatile")
            
            payload = {
                "model": model,
                "messages": analysis_messages,
                "temperature": 0.3,
                "max_tokens": 4000  # Increased from 2000 to allow more detailed responses
            }
            
            print(f"🔄 Making Groq analysis API call...")
            response = requests.post(api_url, headers=headers, json=payload, timeout=CONFIG.get('API_TIMEOUT', 120))
            print(f"📊 Groq Analysis API Response Status: {response.status_code}")
            
            if response.status_code == 200:
                result = response.json()
                ai_response = result["choices"][0]["message"]["content"]
            else:
                raise Exception(f"Groq Analysis API Error {response.status_code}: {response.text}")

        print("✅ Analysis completed successfully")
        
        # Mark file as analyzed in context
        ContextManager.mark_file_analyzed(session_id, filename)
        
        # Add analysis to conversation history
        analysis_summary = f"Completed analysis of {filename}. Key findings and recommendations provided."
        ContextManager.add_message(session_id, "assistant", analysis_summary)
        
        return jsonify({
            "filename": filename,
            "analysis": ai_response,
            "session_id": session_id
        })

    except requests.exceptions.RequestException as e:
        error_msg = f"Analysis API request failed: {str(e)}"
        print(f"❌ Request error: {error_msg}")
        print("🔄 Falling back to basic analysis...")
        
        # Provide fallback analysis when API fails
        fallback_response = get_fallback_analysis(filename)
        return jsonify({
            "filename": filename,
            "analysis": f"⚠️ AI analysis temporarily unavailable due to network issues.\n\n{fallback_response}",
            "fallback": True
        })
        
    except Exception as e:
        print(f"❌ Analysis error: {str(e)}")
        traceback.print_exc()
        
        # Provide fallback analysis for any other errors
        fallback_response = get_fallback_analysis(filename)
        return jsonify({
            "filename": filename,
            "analysis": f"⚠️ Analysis error occurred: {str(e)}\n\n{fallback_response}",
            "fallback": True
        })

@app.route('/api/context', methods=['GET'])
def get_context():
    """Get conversation context and file history"""
    try:
        session_id = request.args.get('session_id', 'default_session')
        session = ContextManager.get_session(session_id)
        
        # Expose last 20 non-system messages for frontend restore
        visible = [m for m in session["conversation_history"] if m.get("role") != "system"][-20:]
        return jsonify({
            "session_id": session_id,
            "conversation_history": visible,
            "uploaded_files": session["uploaded_files"],
            "total_messages": len(session["conversation_history"]),
            "total_files": len(session["uploaded_files"])
        })
        
    except Exception as e:
        print(f"❌ Context retrieval error: {str(e)}")
        return jsonify({"error": f"Context retrieval failed: {str(e)}"}), 500


@app.route('/api/context/sync', methods=['POST'])
def sync_context():
    """Sync messages from frontend to backend session (for chat persistence)"""
    try:
        data = request.get_json() or {}
        session_id = data.get('session_id', 'default_session')
        messages = data.get('messages', [])
        if not isinstance(messages, list):
            return jsonify({"error": "messages must be an array"}), 400
        db_clear_session(session_id)
        session = ContextManager.get_session(session_id)
        session["conversation_history"] = [
            m for m in session["conversation_history"]
            if m.get("role") == "system"
        ]
        for m in messages:
            role = m.get("role")
            content = m.get("content", "")
            if role in ("user", "assistant") and content:
                ContextManager.add_message(session_id, role, content)
        return jsonify({"success": True, "session_id": session_id})
    except Exception as e:
        print(f"❌ Context sync error: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/settings/provider', methods=['POST', 'GET'])
def settings_provider():
    """GET: return active provider. POST {provider}: switch (groq/openai only)."""
    if request.method == 'GET':
        return jsonify({"provider": get_ai_provider()})
    data = request.get_json() or {}
    new_provider = data.get("provider", "").strip().lower()
    if new_provider == "gemini":
        new_provider = "groq"
    valid = {"groq", "openai"}
    if new_provider not in valid:
        return jsonify({"error": f"Invalid provider. Choose from: {', '.join(sorted(valid))}"}), 400
    os.environ["AI_PROVIDER"] = new_provider
    print(f"🔄 AI provider switched to: {new_provider}")
    return jsonify({"success": True, "provider": new_provider})


# Story type options for "write a user story" flow
STORY_TYPE_OPTIONS = ["User Story", "Use Case", "Requirement Story", "Epic Story", "Feature Story"]
STORY_TYPE_DESCRIPTIONS = {
    "User Story": "A natural-language description of a software feature from the end-user's perspective, typically following the format \"As a [user], I want to [perform some task] so that [I can achieve some goal]\"",
    "Use Case": "A written description of how a user will interact with a software system to achieve a specific goal, typically including actors, preconditions, and postconditions",
    "Requirement Story": "A detailed description of a software requirement, including functional and non-functional requirements, acceptance criteria, and constraints or assumptions",
    "Epic Story": "A high-level description of a software feature too large for a single iteration, often broken down into smaller user stories or tasks",
    "Feature Story": "A description of a specific software feature, including functional and non-functional requirements, acceptance criteria, and constraints or assumptions",
}
STORY_CHOICE_MARKER = "[STORY_CHOICE_OFFERED]"

# Email flow options
EMAIL_GENERATE_OPTIONS = ["Generate it for me", "I will write"]
EMAIL_REPHRASE_OPTIONS = ["Yes, rephrase or help", "No"]
EMAIL_SEND_CONFIRM_OPTIONS = ["Yes, send it", "No"]
EMAIL_GENERATE_CHOICE_MARKER = "[EMAIL_GENERATE_CHOICE]"
EMAIL_REPHRASE_CHOICE_MARKER = "[EMAIL_REPHRASE_CHOICE]"
EMAIL_SEND_CONFIRM_MARKER = "[EMAIL_SEND_CONFIRM]"
EMAIL_RECIPIENT_NEEDED_MARKER = "[EMAIL_RECIPIENT_NEEDED]"


def _is_story_write_request(msg):
    """Detect if user is asking to write/create a story (user story, use case, etc.)"""
    if not msg or len(msg) < 10:
        return False
    m = msg.lower().strip()
    triggers = ["write a user story", "write user story", "create a user story", "create user story",
                "write a use case", "write use case", "create use case",
                "write a requirement story", "write requirement story",
                "write an epic story", "write epic story",
                "write a feature story", "write feature story",
                "write a story", "write story", "create a story",
                "help me write a user story", "i want to write a user story", "i need to write a user story"]
    return any(t in m for t in triggers)


def _is_story_type_selection(msg):
    """Check if message is exactly one of the story type options"""
    return msg.strip() in STORY_TYPE_OPTIONS


def _is_email_request(msg):
    """Detect if user wants to write an email"""
    if not msg or len(msg) < 5:
        return False
    m = msg.lower().strip()
    triggers = ["/email", "write an email", "write email", "compose email", "help me write an email", "i need to write an email"]
    return any(t in m for t in triggers)


def _is_email_option_selection(msg, options):
    """Check if message is exactly one of the given options"""
    return msg.strip() in options


def _get_email_flow(session_id):
    """Get or init email flow state for session"""
    session = ContextManager.get_session(session_id)
    if "email_flow" not in session:
        session["email_flow"] = {"active": False, "step": None, "choice": None, "need_help": None, "email_content": None}
    return session["email_flow"]


def _get_last_assistant_content(session_id):
    """Get content of last assistant message in session"""
    session = ContextManager.get_session(session_id)
    for m in reversed(session["conversation_history"]):
        if m.get("role") == "assistant":
            return m.get("content", "")
    return ""


def _extract_email_draft_from_history(session_id):
    """Extract the actual email draft from conversation (the one before Send email? / recipient prompt)."""
    session = ContextManager.get_session(session_id)
    for m in reversed(session["conversation_history"]):
        if m.get("role") != "assistant":
            continue
        raw = m.get("content", "")
        if EMAIL_SEND_CONFIRM_MARKER not in raw:
            continue
        part = raw.split(EMAIL_SEND_CONFIRM_MARKER)[0].strip()
        part = part.split("\n\nSend email?")[0].strip()
        if part and len(part) > 5 and EMAIL_RECIPIENT_NEEDED_MARKER not in part:
            return part
    return None


@app.route('/api/chat/stream', methods=['POST'])
def chat_stream():
    """Streaming version of /api/chat — emits text chunks via SSE."""
    from flask import Response, stream_with_context

    data      = request.get_json() or {}
    user_msg  = data.get("message", "").strip()
    session_id= data.get("session_id", "default_session")

    if not user_msg:
        return jsonify({"error": "No message provided"}), 400

    # Check email flow FIRST - when in recipient_needed, plain email is the recipient, NOT a generic send command
    email_flow = _get_email_flow(session_id)
    if email_flow.get("step") == "recipient_needed" and "@" in user_msg and "." in user_msg:
        import re
        if re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', user_msg):
            recipient = user_msg.strip()
            content = email_flow.get("email_content") or _extract_email_draft_from_history(session_id) or ""
            if not content or EMAIL_RECIPIENT_NEEDED_MARKER in content or "What email address should this be sent to" in content:
                content = "No content"
            body = create_simple_email_body(content, session_id)
            success, msg = send_email(recipient, "Message from Vise-AI", body)
            email_flow["active"] = False
            email_flow["step"] = None
            ContextManager.add_message(session_id, "user", user_msg)
            result_msg = f"Email sent successfully to {recipient}!" if success else f"Failed to send: {msg}"
            ContextManager.add_message(session_id, "assistant", result_msg)
            def _email_sent_stream():
                yield f"data: {json.dumps({'chunk': result_msg})}\n\n"
                yield f"data: {json.dumps({'done': True, 'session_id': session_id})}\n\n"
            return Response(
                stream_with_context(_email_sent_stream()),
                mimetype="text/event-stream",
                headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no", "Connection": "keep-alive"}
            )

    # Handle other email send commands (e.g. "send X to user@mail.com") before hitting AI
    email_command = detect_email_command(user_msg)
    if email_command['is_email_command']:
        if email_command.get('send_previous_response'):
            session = ContextManager.get_session(session_id)
            last_ai_response = None
            for msg in reversed(session["conversation_history"]):
                if msg["role"] == "assistant":
                    last_ai_response = msg["content"]
                    break
            message_content = last_ai_response or "No previous AI response found to send."
            subject = "AI Response from Vise-AI"
        else:
            message_content = email_command['message_content']
            subject = "Message from Vise-AI User"
        body = create_simple_email_body(message_content, session_id)
        success, msg = send_email(email_command['recipient_email'], subject, body)
        if success:
            ai_response = f"Email sent successfully to {email_command['recipient_email']}!"
        else:
            ai_response = f"Failed to send email: {msg}"
        ContextManager.add_message(session_id, "assistant", ai_response)

        def _email_done_stream():
            yield f"data: {json.dumps({'chunk': ai_response})}\n\n"
            yield f"data: {json.dumps({'done': True, 'session_id': session_id})}\n\n"
        return Response(
            stream_with_context(_email_done_stream()),
            mimetype="text/event-stream",
            headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no", "Connection": "keep-alive"}
        )

    last_assistant = _get_last_assistant_content(session_id)
    is_story_selection = _is_story_type_selection(user_msg) and STORY_CHOICE_MARKER in last_assistant
    email_flow = _get_email_flow(session_id)

    # ─── Email flow handling ─────────────────────────────────────────────────
    if _is_email_request(user_msg) and not email_flow.get("active"):
        ContextManager.add_message(session_id, "user", user_msg)
        email_flow["active"] = True
        email_flow["step"] = "generate_choice"
        choice_msg = "Should I generate the email for you?"
        stored_msg = f"{choice_msg}\n\n{EMAIL_GENERATE_CHOICE_MARKER}\n\nSelect one of the options above."
        ContextManager.add_message(session_id, "assistant", stored_msg)

        def _email_choice_stream():
            yield f"data: {json.dumps({'email_generate_choice': True, 'message': choice_msg, 'options': EMAIL_GENERATE_OPTIONS})}\n\n"
            yield f"data: {json.dumps({'done': True, 'session_id': session_id})}\n\n"
        return Response(
            stream_with_context(_email_choice_stream()),
            mimetype="text/event-stream",
            headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no", "Connection": "keep-alive"}
        )

    if _is_email_option_selection(user_msg, EMAIL_GENERATE_OPTIONS) and EMAIL_GENERATE_CHOICE_MARKER in last_assistant:
        ContextManager.add_message(session_id, "user", user_msg)
        email_flow["choice"] = user_msg.strip()
        if user_msg.strip() == "Generate it for me":
            email_flow["step"] = "waiting_generate_content"
            prompt_msg = "Please share the content of the email (what you want to say, who it's for, purpose, etc.). I won't write a blind email."
            stored_msg = prompt_msg
            ContextManager.add_message(session_id, "assistant", stored_msg)

            def _email_content_prompt_stream():
                yield f"data: {json.dumps({'chunk': stored_msg})}\n\n"
                yield f"data: {json.dumps({'done': True, 'session_id': session_id})}\n\n"
            return Response(
                stream_with_context(_email_content_prompt_stream()),
                mimetype="text/event-stream",
                headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no", "Connection": "keep-alive"}
            )
        else:
            email_flow["step"] = "rephrase_choice"
            choice_msg = "Should I rephrase or help you with your email?"
            stored_msg = f"{choice_msg}\n\n{EMAIL_REPHRASE_CHOICE_MARKER}\n\nSelect one of the options above."
            ContextManager.add_message(session_id, "assistant", stored_msg)

            def _email_rephrase_stream():
                yield f"data: {json.dumps({'email_rephrase_choice': True, 'message': choice_msg, 'options': EMAIL_REPHRASE_OPTIONS})}\n\n"
                yield f"data: {json.dumps({'done': True, 'session_id': session_id})}\n\n"
            return Response(
                stream_with_context(_email_rephrase_stream()),
                mimetype="text/event-stream",
                headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no", "Connection": "keep-alive"}
            )

    if _is_email_option_selection(user_msg, EMAIL_REPHRASE_OPTIONS) and EMAIL_REPHRASE_CHOICE_MARKER in last_assistant:
        ContextManager.add_message(session_id, "user", user_msg)
        email_flow["need_help"] = user_msg.strip().startswith("Yes")
        email_flow["step"] = "waiting_content"
        hint = "Type your email content below. I'll help rephrase it." if email_flow["need_help"] else "Type your email content below."
        stored_msg = f"{hint}\n\nWhen you're done, you'll get an option to send."
        ContextManager.add_message(session_id, "assistant", stored_msg)

        def _email_wait_stream():
            yield f"data: {json.dumps({'chunk': stored_msg})}\n\n"
            yield f"data: {json.dumps({'done': True, 'session_id': session_id})}\n\n"
        return Response(
            stream_with_context(_email_wait_stream()),
            mimetype="text/event-stream",
            headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no", "Connection": "keep-alive"}
        )

    # Email send confirm: user clicked "Yes, send it" -> need recipient
    if _is_email_option_selection(user_msg, EMAIL_SEND_CONFIRM_OPTIONS) and EMAIL_SEND_CONFIRM_MARKER in last_assistant:
        ContextManager.add_message(session_id, "user", user_msg)
        if user_msg.strip() == "Yes, send it":
            email_flow["step"] = "recipient_needed"
            draft = _extract_email_draft_from_history(session_id) or email_flow.get("email_content")
            email_flow["email_content"] = draft
            prompt_msg = "What email address should this be sent to?"
            stored_msg = f"{prompt_msg}\n\n{EMAIL_RECIPIENT_NEEDED_MARKER}"
            ContextManager.add_message(session_id, "assistant", stored_msg)

            def _email_recipient_stream():
                yield f"data: {json.dumps({'email_recipient_needed': True, 'message': prompt_msg})}\n\n"
                yield f"data: {json.dumps({'done': True, 'session_id': session_id})}\n\n"
            return Response(
                stream_with_context(_email_recipient_stream()),
                mimetype="text/event-stream",
                headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no", "Connection": "keep-alive"}
            )
        else:
            email_flow["active"] = False
            email_flow["step"] = None
            stored_msg = "Email not sent. You can start a new email anytime with /email"
            ContextManager.add_message(session_id, "assistant", stored_msg)
            def _email_cancel_stream():
                yield f"data: {json.dumps({'chunk': stored_msg})}\n\n"
                yield f"data: {json.dumps({'done': True, 'session_id': session_id})}\n\n"
            return Response(
                stream_with_context(_email_cancel_stream()),
                mimetype="text/event-stream",
                headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no", "Connection": "keep-alive"}
            )

    # Email: user shared content for "Generate it for me" -> now generate
    if email_flow.get("step") == "waiting_generate_content":
        ContextManager.add_message(session_id, "user", user_msg)
        email_flow["step"] = "generating"
        email_flow["generate_request"] = user_msg
        # Fall through to AI with inject

    # Email: user wrote content without help - show send confirm directly
    if email_flow.get("step") == "waiting_content" and not email_flow.get("need_help"):
        ContextManager.add_message(session_id, "user", user_msg)
        email_flow["email_content"] = user_msg
        email_flow["step"] = "send_confirm"
        confirm_msg = "Send email?"
        stored_msg = f"{confirm_msg}\n\n{EMAIL_SEND_CONFIRM_MARKER}\n\nSelect one of the options above."
        ContextManager.add_message(session_id, "assistant", stored_msg)

        def _email_confirm_stream():
            yield f"data: {json.dumps({'email_send_confirm': True, 'message': confirm_msg, 'options': EMAIL_SEND_CONFIRM_OPTIONS})}\n\n"
            yield f"data: {json.dumps({'done': True, 'session_id': session_id})}\n\n"
        return Response(
            stream_with_context(_email_confirm_stream()),
            mimetype="text/event-stream",
            headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no", "Connection": "keep-alive"}
        )

    if email_flow.get("step") != "generating":
        cleaned = tonnify_preprocess(user_msg)
        ContextManager.add_message(session_id, "user", cleaned if cleaned else user_msg)

    # Case 1: User asked to write a story (first time) -> show choice, no AI call
    if _is_story_write_request(user_msg) and not is_story_selection:
        choice_msg = "Which type of story would you like to write?"
        stored_msg = f"{choice_msg}\n\n{STORY_CHOICE_MARKER}\n\nSelect one of the options above."
        ContextManager.add_message(session_id, "assistant", stored_msg)

        def _story_choice_stream():
            yield f"data: {json.dumps({'story_type_choice': True, 'message': choice_msg, 'options': STORY_TYPE_OPTIONS})}\n\n"
            yield f"data: {json.dumps({'done': True, 'session_id': session_id})}\n\n"

        return Response(
            stream_with_context(_story_choice_stream()),
            mimetype="text/event-stream",
            headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no", "Connection": "keep-alive"}
        )

    messages = ContextManager.get_conversation_context(session_id, include_files=True)
    provider = get_ai_provider()

    # Case 2: User selected a story type -> inject prompt for that type
    if is_story_selection:
        desc = STORY_TYPE_DESCRIPTIONS.get(user_msg.strip(), "")
        inject = {
            "role": "system",
            "content": f"The user previously asked to write a story and selected: **{user_msg}**. "
            f"Generate a comprehensive {user_msg} following this definition: {desc}. "
            f"Include: 1) Format/structure explanation, 2) A complete example, 3) Best practices."
        }
        messages.insert(1, inject)

    # Case 2b: Email flow - generate draft (user shared content, don't write blind)
    if email_flow.get("step") == "generating":
        content_req = email_flow.get("generate_request", user_msg) or user_msg
        reply_hint = ""
        if any(p in content_req.lower() for p in ("reply", "repl", "respond", "answer", "responding")):
            reply_hint = (
                "IMPORTANT: The user said they need to REPLY to an email. They ARE the one sending the reply. "
                "Write the email AS the user's reply - from their perspective, as if they are responding to the other person. "
            )
        inject = {
            "role": "system",
            "content": f"The user wants you to generate an email draft based on their request. "
            f"Their request: \"{content_req}\" "
            f"{reply_hint}"
            f"Create a professional, well-formatted email that addresses what they asked for. "
            f"Output ONLY the email body content - exactly what will be sent. No meta-commentary, no placeholders like [Your Name] unless the user explicitly asked for them."
        }
        messages.insert(1, inject)

    # Case 2c: Email flow - rephrase user's email
    if email_flow.get("step") == "waiting_content" and email_flow.get("need_help"):
        inject = {
            "role": "system",
            "content": f"The user shared their email draft and wants you to rephrase or improve it. "
            f"Their draft: \"{user_msg}\" "
            f"Provide a polished, professional version. Output ONLY the improved email body, no extra commentary."
        }
        messages.insert(1, inject)

    def _generate():
        full_response = []
        try:
            if provider == "openai":
                import openai as _oai
                _oai_client = _oai.OpenAI(api_key=CONFIG["OPENAI_API_KEY"])
                stream = _oai_client.chat.completions.create(
                    model=CONFIG.get("OPENAI_MODEL", "o4-mini-2025-04-16"),
                    messages=messages, temperature=0.6, max_tokens=8192, stream=True
                )
                for chunk in stream:
                    delta = chunk.choices[0].delta.content or ""
                    if delta:
                        full_response.append(delta)
                        yield f"data: {json.dumps({'chunk': delta})}\n\n"

            else:  # groq
                from groq import Groq as _Groq
                gclient = _Groq(api_key=CONFIG["GROQ_API_KEY"])
                stream = gclient.chat.completions.create(
                    model=CONFIG.get("GROQ_MODEL", "llama-3.3-70b-versatile"),
                    messages=messages, temperature=0.6, max_tokens=8192, stream=True
                )
                for chunk in stream:
                    delta = chunk.choices[0].delta.content or ""
                    if delta:
                        full_response.append(delta)
                        yield f"data: {json.dumps({'chunk': delta})}\n\n"

            ai_response = "".join(full_response)

            # Email flow: after generating or rephrasing, show send confirm
            show_send_confirm = (
                email_flow.get("step") == "generating" or
                (email_flow.get("step") == "waiting_content" and email_flow.get("need_help"))
            )
            if show_send_confirm:
                email_flow["email_content"] = ai_response
                email_flow["step"] = "send_confirm"
                confirm_msg = "Send email?"
                full_msg = f"{ai_response}\n\n{confirm_msg}\n\n{EMAIL_SEND_CONFIRM_MARKER}"
                ContextManager.add_message(session_id, "assistant", full_msg)
                yield f"data: {json.dumps({'email_send_confirm': True, 'message': confirm_msg, 'options': EMAIL_SEND_CONFIRM_OPTIONS})}\n\n"
            else:
                ContextManager.add_message(session_id, "assistant", ai_response)

            yield f"data: {json.dumps({'done': True, 'session_id': session_id})}\n\n"

        except Exception as e:
            print(f"❌ Streaming error: {e}")
            traceback.print_exc()
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return Response(
        stream_with_context(_generate()),
        mimetype="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive"
        }
    )


@app.route('/api/chat/history', methods=['GET'])
def chat_history_list():
    """Return all persisted sessions with message counts for the History view."""
    sessions = db_all_sessions()
    return jsonify({"sessions": sessions, "total": len(sessions)})


@app.route('/api/chat/history/<session_id>', methods=['GET'])
def chat_history_session(session_id):
    """Return full message history for one session from DB."""
    limit = int(request.args.get("limit", 100))
    messages = db_load_session(session_id, limit=limit)
    return jsonify({"session_id": session_id, "messages": messages, "total": len(messages)})


@app.route('/api/summarize-files', methods=['POST'])
def summarize_files():
    """Summarise all uploaded files in a session so the AI has multi-file context."""
    data       = request.get_json() or {}
    session_id = data.get("session_id", "default_session")
    session    = ContextManager.get_session(session_id)
    files      = session.get("uploaded_files", [])

    if not files:
        return jsonify({"error": "No uploaded files found in session"}), 400

    summaries = []
    for f in files:
        fpath = f.get("path", "")
        fname = f.get("filename", "unknown")
        try:
            if f.get("type") == "image":
                summaries.append(f"📷 {fname} — image file (not summarised)")
                continue
            ext = os.path.splitext(fname)[1].lower()
            if ext == ".pdf":
                content = extract_pdf_content(fpath)
            elif ext in (".docx", ".doc"):
                content = extract_docx_content(fpath)
            else:
                with open(fpath, "r", encoding="utf-8", errors="replace") as fh:
                    content = fh.read(10000)
            summaries.append(f"📄 **{fname}**:\n{content[:500]}…")
        except Exception as e:
            summaries.append(f"📄 {fname} — could not read: {e}")

    combined = "\n\n".join(summaries)
    # Inject this as system context into the session
    session["conversation_history"].append({
        "role": "system",
        "content": f"[Multi-file context loaded]\n{combined}",
        "timestamp": time.time()
    })
    return jsonify({"success": True, "files_loaded": len(files), "preview": combined[:300]})


@app.route('/api/key-status', methods=['GET'])
def get_key_status():
    """Return which API keys are configured (masked, never expose full keys)."""
    def _masked(key: str) -> str:
        return f"{key[:8]}…{key[-4:]}" if key and len(key) > 12 else ""

    groq_key   = CONFIG.get("GROQ_API_KEY", "")
    openai_key = CONFIG.get("OPENAI_API_KEY", "")

    return jsonify({
        "groq":   {"set": bool(groq_key),   "preview": _masked(groq_key)},
        "openai": {"set": bool(openai_key), "preview": _masked(openai_key)},
        "active_provider": get_ai_provider(),
    })


@app.route('/api/clear-context', methods=['POST'])
def clear_context():
    """Clear conversation context for a session"""
    try:
        data = request.get_json()
        session_id = data.get("session_id", "default_session")
        
        if session_id in user_sessions:
            del user_sessions[session_id]
        db_clear_session(session_id)
        print(f"🗑️ Cleared context for session: {session_id}")
        
        return jsonify({
            "success": True,
            "message": f"Context cleared for session: {session_id}"
        })
        
    except Exception as e:
        print(f"❌ Context clearing error: {str(e)}")
        return jsonify({"error": f"Context clearing failed: {str(e)}"}), 500

@app.route('/api/email-analysis', methods=['POST'])
def email_analysis():
    """Send analysis report via email"""
    try:
        data = request.get_json() or {}
        filename = data.get("filename")
        analysis = data.get("analysis")
        recipient = (data.get("recipient") or CONFIG.get("DEFAULT_RECIPIENT") or "").strip()
        session_id = data.get("session_id", "default_session")

        if not filename or not analysis:
            return jsonify({"error": "Filename and analysis are required"}), 400

        if not recipient:
            return jsonify({
                "error": "Recipient email is required. Set DEFAULT_RECIPIENT in .env or provide recipient when sending."
            }), 400

        if not CONFIG.get('EMAIL_ENABLED', False):
            return jsonify({"error": "Email functionality is disabled"}), 400

        # Create email content
        subject = f"📊 Vise-AI Analysis Report: {filename}"
        session_info = f"Session: {session_id} | Files analyzed: 1"
        body = create_analysis_email_body(filename, analysis, session_info)
        
        # Send email
        success, message = send_email(recipient, subject, body)
        
        if success:
            return jsonify({
                "success": True,
                "message": f"Analysis report sent to {recipient}",
                "recipient": recipient
            })
        else:
            return jsonify({"error": message}), 500
            
    except Exception as e:
        print(f"❌ Email analysis error: {str(e)}")
        return jsonify({"error": f"Failed to send email: {str(e)}"}), 500

@app.route('/api/email-notification', methods=['POST'])
def email_notification():
    """Send notification email"""
    try:
        data = request.get_json()
        title = data.get("title", "Vise-AI Notification")
        message = data.get("message")
        details = data.get("details")
        recipient = data.get("recipient", CONFIG.get("DEFAULT_RECIPIENT"))
        
        if not message:
            return jsonify({"error": "Message is required"}), 400
        
        if not CONFIG.get('EMAIL_ENABLED', False):
            return jsonify({"error": "Email functionality is disabled"}), 400
        
        # Create email content
        subject = f"🔔 {title}"
        body = create_notification_email_body(title, message, details)
        
        # Send email
        success, email_message = send_email(recipient, subject, body)
        
        if success:
            return jsonify({
                "success": True,
                "message": f"Notification sent to {recipient}",
                "recipient": recipient
            })
        else:
            return jsonify({"error": email_message}), 500
            
    except Exception as e:
        print(f"❌ Email notification error: {str(e)}")
        return jsonify({"error": f"Failed to send notification: {str(e)}"}), 500

@app.route('/api/email-config', methods=['GET'])
def get_email_config():
    """Get email configuration status"""
    try:
        return jsonify({
            "email_enabled": CONFIG.get('EMAIL_ENABLED', False),
            "smtp_server": CONFIG.get('SMTP_SERVER', ''),
            "sender_email": CONFIG.get('SENDER_EMAIL', ''),
            "owner_email": CONFIG.get('OWNER_EMAIL', ''),
            "default_recipient": CONFIG.get('DEFAULT_RECIPIENT', ''),
            "smtp_configured": all([
                CONFIG.get('SMTP_SERVER'),
                CONFIG.get('SENDER_EMAIL'),
                CONFIG.get('APP_PASSWORD')
            ])
        })
    except Exception as e:
        print(f"❌ Email config error: {str(e)}")
        return jsonify({"error": f"Failed to get email config: {str(e)}"}), 500

@app.route('/api/test-email', methods=['POST'])
def test_email():
    """Test email functionality"""
    try:
        data = request.get_json() or {}
        recipient = (data.get("recipient") or CONFIG.get("DEFAULT_RECIPIENT") or "").strip()

        if not recipient:
            return jsonify({
                "error": "Recipient email is required. Set DEFAULT_RECIPIENT in .env or provide recipient in the request."
            }), 400

        if not CONFIG.get('EMAIL_ENABLED', False):
            return jsonify({"error": "Email functionality is disabled"}), 400

        # Send test email
        subject = "🧪 Vise-AI Email Test"
        message = "This is a test email from your Vise-AI system to verify email functionality is working correctly."
        details = f"Sent from: {CONFIG.get('SENDER_EMAIL')}<br>SMTP Server: {CONFIG.get('SMTP_SERVER')}<br>Test Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        
        body = create_notification_email_body("Email Test", message, details)
        success, email_message = send_email(recipient, subject, body)
        
        if success:
            return jsonify({
                "success": True,
                "message": f"Test email sent successfully to {recipient}",
                "recipient": recipient
            })
        else:
            return jsonify({"error": email_message}), 500
            
    except Exception as e:
        print(f"❌ Test email error: {str(e)}")
        return jsonify({"error": f"Failed to send test email: {str(e)}"}), 500

@app.route('/api/generate-document', methods=['POST'])
def generate_document():
    """Generate document in specified format"""
    try:
        data = request.get_json()
        content = data.get('content')
        format_type = data.get('format', 'pdf').lower()
        filename = data.get('filename', 'vise-ai-document')
        
        if not content:
            return jsonify({"error": "Content is required"}), 400
        
        print(f"📄 Generating {format_type.upper()} document...")
        print(f"📊 Content length: {len(content)} characters")
        
        # Generate document based on format
        if format_type == 'pdf':
            file_content, error = generate_pdf_document(content, f"{filename}.pdf")
            mime_type = 'application/pdf'
            extension = '.pdf'
        elif format_type == 'docx':
            file_content, error = generate_docx_document(content, f"{filename}.docx")
            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            extension = '.docx'
        elif format_type == 'excel' or format_type == 'xlsx':
            file_content, error = generate_excel_document(content, f"{filename}.xlsx")
            mime_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            extension = '.xlsx'
        elif format_type == 'txt':
            file_content, error = generate_txt_document(content, f"{filename}.txt")
            mime_type = 'text/plain'
            extension = '.txt'
        else:
            return jsonify({"error": f"Unsupported format: {format_type}"}), 400
        
        if error:
            print(f"❌ Document generation error: {error}")
            return jsonify({"error": error}), 500
        
        # Generate unique file ID
        file_id = str(uuid.uuid4())
        
        # Create downloads directory
        os.makedirs('downloads', exist_ok=True)
        file_path = os.path.join('downloads', f"{file_id}{extension}")
        
        # Save file (handle text vs binary content)
        if format_type == 'txt':
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(file_content)
        else:
            with open(file_path, 'wb') as f:
                f.write(file_content)
        
        print(f"✅ Document generated: {file_path}")
        
        return jsonify({
            "success": True,
            "file_id": file_id,
            "filename": f"{filename}{extension}",
            "format": format_type,
            "size": len(file_content),
            "message": f"{format_type.upper()} document generated successfully"
        })
        
    except Exception as e:
        print(f"❌ Document generation error: {str(e)}")
        return jsonify({"error": f"Document generation failed: {str(e)}"}), 500

@app.route('/api/download/<file_id>', methods=['GET'])
def download_file(file_id):
    """Download generated file"""
    try:
        print(f"📥 Download request for file ID: {file_id}")
        
        # Find the file in downloads directory
        downloads_dir = 'downloads'
        if not os.path.exists(downloads_dir):
            return jsonify({"error": "Downloads directory not found"}), 404
        
        # Look for file with this ID
        for filename in os.listdir(downloads_dir):
            if filename.startswith(file_id):
                file_path = os.path.join(downloads_dir, filename)
                
                # Determine MIME type based on extension
                if filename.endswith('.pdf'):
                    mime_type = 'application/pdf'
                elif filename.endswith('.docx'):
                    mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                elif filename.endswith('.xlsx'):
                    mime_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                elif filename.endswith('.txt'):
                    mime_type = 'text/plain'
                else:
                    mime_type = 'application/octet-stream'
                
                print(f"✅ File found: {filename}")
                
                return send_file(
                    file_path,
                    as_attachment=True,
                    download_name=filename,
                    mimetype=mime_type
                )
        
        print(f"❌ File not found: {file_id}")
        return jsonify({"error": "File not found"}), 404
        
    except Exception as e:
        print(f"❌ Download error: {str(e)}")
        return jsonify({"error": f"Download failed: {str(e)}"}), 500

@app.route('/api/documents', methods=['GET'])
def get_documents():
    """Get all uploaded documents for a session"""
    try:
        session_id = request.args.get('session_id', 'default_session')
        session = ContextManager.get_session(session_id)
        
        # Get uploaded files with additional details
        documents = []
        for file_info in session["uploaded_files"]:
            file_path = file_info["path"]
            file_exists = os.path.exists(file_path)
            
            # Get file stats if file exists
            file_stats = {}
            if file_exists:
                try:
                    stat = os.stat(file_path)
                    file_stats = {
                        "size_bytes": stat.st_size,
                        "size_mb": round(stat.st_size / (1024 * 1024), 2),
                        "created": stat.st_ctime,
                        "modified": stat.st_mtime
                    }
                except:
                    pass
            
            document = {
                "filename": file_info["filename"],
                "file_id": file_info.get("fileId", ""),
                "path": file_path,
                "size": file_info["size"],
                "type": file_info.get("type", "unknown"),
                "extension": file_info.get("extension", ""),
                "upload_time": file_info["upload_time"],
                "analyzed": file_info["analyzed"],
                "exists": file_exists,
                "stats": file_stats
            }
            documents.append(document)
        
        # Sort by upload time (newest first)
        documents.sort(key=lambda x: x["upload_time"], reverse=True)
        
        return jsonify({
            "success": True,
            "session_id": session_id,
            "documents": documents,
            "total_count": len(documents),
            "analyzed_count": sum(1 for doc in documents if doc["analyzed"]),
            "total_size_mb": sum(doc.get("stats", {}).get("size_mb", 0) for doc in documents)
        })
        
    except Exception as e:
        print(f"❌ Documents retrieval error: {str(e)}")
        return jsonify({"error": f"Failed to get documents: {str(e)}"}), 500

@app.route('/api/documents/<file_id>', methods=['GET'])
def get_document_details(file_id):
    """Get detailed information about a specific document"""
    try:
        session_id = request.args.get('session_id', 'default_session')
        session = ContextManager.get_session(session_id)
        
        # Find the document
        document = None
        for file_info in session["uploaded_files"]:
            if file_info.get("fileId") == file_id:
                document = file_info
                break
        
        if not document:
            return jsonify({"error": "Document not found"}), 404
        
        file_path = document["path"]
        file_exists = os.path.exists(file_path)
        
        # Get detailed file information
        file_details = {
            "filename": document["filename"],
            "file_id": file_id,
            "path": file_path,
            "size": document["size"],
            "type": document.get("type", "unknown"),
            "extension": document.get("extension", ""),
            "upload_time": document["upload_time"],
            "analyzed": document["analyzed"],
            "exists": file_exists
        }
        
        if file_exists:
            try:
                stat = os.stat(file_path)
                file_details.update({
                    "size_bytes": stat.st_size,
                    "size_mb": round(stat.st_size / (1024 * 1024), 2),
                    "created": stat.st_ctime,
                    "modified": stat.st_mtime,
                    "created_readable": datetime.fromtimestamp(stat.st_ctime).strftime('%Y-%m-%d %H:%M:%S'),
                    "modified_readable": datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S')
                })
            except Exception as e:
                file_details["error"] = f"Could not read file stats: {str(e)}"
        
        return jsonify({
            "success": True,
            "document": file_details
        })
        
    except Exception as e:
        print(f"❌ Document details error: {str(e)}")
        return jsonify({"error": f"Failed to get document details: {str(e)}"}), 500

@app.route('/api/documents/<file_id>/content', methods=['GET'])
def get_document_content(file_id):
    """Get file content for viewing (text) or serve file (images/binary)"""
    from flask import send_file
    try:
        session_id = request.args.get('session_id', 'default_session')
        session = ContextManager.get_session(session_id)
        document = None
        for f in session["uploaded_files"]:
            if f.get("fileId") == file_id:
                document = f
                break
        if not document:
            return jsonify({"error": "Document not found"}), 404
        path = document.get("path")
        if not path or not os.path.exists(path):
            return jsonify({"error": "File not found on disk"}), 404
        ext = (document.get("extension") or "").lower()
        if ext in ('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp'):
            mime = {'jpg':'jpeg','jpeg':'jpeg','png':'png','gif':'gif','bmp':'bmp','webp':'webp'}.get(ext[1:], 'jpeg')
            return send_file(path, mimetype=f'image/{mime}')
        if ext in ('.txt', '.md', '.csv'):
            with open(path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
            return jsonify({"content": content, "type": "text", "filename": document.get("filename")})
        if ext in ('.pdf', '.docx'):
            return jsonify({"type": "binary", "filename": document.get("filename"), "message": "Use Download or Analyze for this file type"})
        try:
            with open(path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
            return jsonify({"content": content, "type": "text", "filename": document.get("filename")})
        except Exception:
            return jsonify({"type": "binary", "filename": document.get("filename"), "message": "Binary file - use Analyze"})
    except Exception as e:
        print(f"❌ Document content error: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/documents/<file_id>/delete', methods=['DELETE'])
def delete_document(file_id):
    """Delete a specific document"""
    try:
        session_id = request.args.get('session_id', 'default_session')
        session = ContextManager.get_session(session_id)
        
        # Find and remove the document from session
        document = None
        for i, file_info in enumerate(session["uploaded_files"]):
            if file_info.get("fileId") == file_id:
                document = file_info
                session["uploaded_files"].pop(i)
                break
        
        if not document:
            return jsonify({"error": "Document not found in session"}), 404
        
        # Delete the physical file
        file_path = document["path"]
        file_deleted = False
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
                file_deleted = True
                print(f"🗑️ Deleted file: {file_path}")
            except Exception as e:
                print(f"❌ Failed to delete file {file_path}: {str(e)}")
        
        return jsonify({
            "success": True,
            "message": f"Document '{document['filename']}' deleted successfully",
            "file_deleted": file_deleted,
            "session_id": session_id
        })
        
    except Exception as e:
        print(f"❌ Document deletion error: {str(e)}")
        return jsonify({"error": f"Failed to delete document: {str(e)}"}), 500

@app.route('/api/document-formats', methods=['GET'])
def get_document_formats():
    """Get available document formats and their status"""
    try:
        formats = {
            'pdf': {
                'available': REPORTLAB_AVAILABLE,
                'description': 'Portable Document Format',
                'extension': '.pdf'
            },
            'docx': {
                'available': DOCX_AVAILABLE,
                'description': 'Microsoft Word Document',
                'extension': '.docx'
            },
            'excel': {
                'available': OPENPYXL_AVAILABLE,
                'description': 'Microsoft Excel Spreadsheet',
                'extension': '.xlsx'
            },
            'txt': {
                'available': True,
                'description': 'Plain Text Document',
                'extension': '.txt'
            }
        }
        
        return jsonify({
            "formats": formats,
            "total_available": sum(1 for f in formats.values() if f['available'])
        })
        
    except Exception as e:
        print(f"❌ Formats error: {str(e)}")
        return jsonify({"error": f"Failed to get formats: {str(e)}"}), 500

# Calendar Events API Endpoints
@app.route('/api/calendar-events', methods=['GET'])
def get_calendar_events():
    """Get all calendar events with optional filters"""
    try:
        events = load_calendar_events()
        
        # Apply filters
        date_filter = request.args.get('date')
        email_filter = request.args.get('email')
        event_type_filter = request.args.get('event_type')
        
        if date_filter:
            events = [e for e in events if e["date"] == date_filter]
        if email_filter:
            events = [e for e in events if e["email"] == email_filter]
        if event_type_filter:
            events = [e for e in events if e.get("event_type") == event_type_filter]
        
        # Sort by date
        events.sort(key=lambda x: x["date"])
        
        return jsonify({
            "success": True,
            "events": events,
            "total_count": len(events)
        })
        
    except Exception as e:
        print(f"❌ Calendar events retrieval error: {str(e)}")
        return jsonify({"error": f"Failed to get calendar events: {str(e)}"}), 500

@app.route('/api/calendar-events/today', methods=['GET'])
def get_today_events():
    """Get events for today"""
    try:
        today = date.today()
        events = get_events_for_date(today)
        
        return jsonify({
            "success": True,
            "date": today.strftime("%Y-%m-%d"),
            "events": events,
            "count": len(events)
        })
        
    except Exception as e:
        print(f"❌ Today events error: {str(e)}")
        return jsonify({"error": f"Failed to get today events: {str(e)}"}), 500

@app.route('/api/calendar-events', methods=['POST'])
def create_calendar_event_endpoint():
    """Create a new calendar event"""
    try:
        data = request.get_json()
        
        event_date = data.get("date")
        recipient_email = data.get("email")
        image_path = data.get("image_path")
        event_type = data.get("event_type", "anniversary")
        message = data.get("message", "")
        
        if not event_date or not recipient_email:
            return jsonify({"error": "Date and email are required"}), 400
        
        # Validate date format
        try:
            datetime.strptime(event_date, "%Y-%m-%d")
        except ValueError:
            return jsonify({"error": "Invalid date format. Use YYYY-MM-DD"}), 400
        
        event = create_calendar_event(
            event_date,
            recipient_email,
            image_path or "",
            event_type,
            message
        )
        
        return jsonify({
            "success": True,
            "event": event,
            "message": f"Calendar event created for {event_date}"
        }), 201
        
    except Exception as e:
        print(f"❌ Calendar event creation error: {str(e)}")
        return jsonify({"error": f"Failed to create calendar event: {str(e)}"}), 500

@app.route('/api/calendar-events/<event_id>', methods=['GET'])
def get_calendar_event(event_id):
    """Get a specific calendar event"""
    try:
        events = load_calendar_events()
        event = next((e for e in events if e["id"] == event_id), None)
        
        if not event:
            return jsonify({"error": "Event not found"}), 404
        
        return jsonify({
            "success": True,
            "event": event
        })
        
    except Exception as e:
        print(f"❌ Calendar event retrieval error: {str(e)}")
        return jsonify({"error": f"Failed to get calendar event: {str(e)}"}), 500

@app.route('/api/calendar-events/<event_id>', methods=['PUT'])
def update_calendar_event(event_id):
    """Update a calendar event"""
    try:
        events = load_calendar_events()
        event = next((e for e in events if e["id"] == event_id), None)
        
        if not event:
            return jsonify({"error": "Event not found"}), 404
        
        data = request.get_json()
        
        # Update fields
        if "date" in data:
            event["date"] = data["date"]
        if "email" in data:
            event["email"] = data["email"]
        if "image_path" in data:
            event["image_path"] = data["image_path"]
        if "event_type" in data:
            event["event_type"] = data["event_type"]
        if "message" in data:
            event["message"] = data["message"]
        
        save_calendar_events(events)
        
        return jsonify({
            "success": True,
            "event": event,
            "message": "Calendar event updated successfully"
        })
        
    except Exception as e:
        print(f"❌ Calendar event update error: {str(e)}")
        return jsonify({"error": f"Failed to update calendar event: {str(e)}"}), 500

@app.route('/api/calendar-events/<event_id>', methods=['DELETE'])
def delete_calendar_event(event_id):
    """Delete a calendar event"""
    try:
        events = load_calendar_events()
        event = next((e for e in events if e["id"] == event_id), None)
        
        if not event:
            return jsonify({"error": "Event not found"}), 404
        
        # Optionally delete image file
        if event.get("image_path") and os.path.exists(event["image_path"]):
            try:
                os.remove(event["image_path"])
            except:
                pass
        
        events.remove(event)
        save_calendar_events(events)
        
        return jsonify({
            "success": True,
            "message": "Calendar event deleted successfully"
        })
        
    except Exception as e:
        print(f"❌ Calendar event deletion error: {str(e)}")
        return jsonify({"error": f"Failed to delete calendar event: {str(e)}"}), 500

@app.route('/api/calendar-events/upload-image', methods=['POST'])
def upload_calendar_image():
    """Upload an image for calendar events"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        # Validate file type
        allowed_extensions = {'.jpg', '.jpeg', '.png', '.gif'}
        file_extension = os.path.splitext(file.filename)[1].lower()
        
        if file_extension not in allowed_extensions:
            return jsonify({"error": "Invalid file type. Allowed: JPG, PNG, GIF"}), 400
        
        # Save file
        ensure_calendar_directories()
        file_id = str(uuid.uuid4())
        filename = f"{file_id}{file_extension}"
        file_path = os.path.join(CALENDAR_IMAGES_DIR, filename)
        
        file.save(file_path)
        
        return jsonify({
            "success": True,
            "image_path": file_path,
            "filename": filename,
            "message": "Image uploaded successfully"
        })
        
    except Exception as e:
        print(f"❌ Calendar image upload error: {str(e)}")
        return jsonify({"error": f"Failed to upload image: {str(e)}"}), 500

@app.route('/api/calendar-events/send-now/<event_id>', methods=['POST'])
def send_calendar_event_now(event_id):
    """Manually trigger email send for a calendar event"""
    try:
        events = load_calendar_events()
        event = next((e for e in events if e["id"] == event_id), None)
        
        if not event:
            return jsonify({"error": "Event not found"}), 404
        
        if event.get("sent", False):
            return jsonify({"error": "Event already sent"}), 400
        
        success = send_calendar_event_email(event)
        
        if success:
            return jsonify({
                "success": True,
                "message": f"Email sent successfully to {event['email']}"
            })
        else:
            return jsonify({
                "success": False,
                "error": "Failed to send email"
            }), 500
        
    except Exception as e:
        print(f"❌ Calendar event send error: {str(e)}")
        return jsonify({"error": f"Failed to send calendar event: {str(e)}"}), 500

if __name__ == '__main__':
    print("Starting Vise-AI Flask Server...")
    p = get_ai_provider()
    names = {"groq": "Groq", "openai": "OpenAI"}
    print(f"AI Provider: {names.get(p, p)}")
    if p == 'openai':
        key = CONFIG.get("OPENAI_API_KEY", "")
        print(f"OpenAI model: {CONFIG.get('OPENAI_MODEL', 'o4-mini-2025-04-16')}, key: ...{key[-4:] if key and len(key) > 4 else 'NOT SET'}")
    
    # Start calendar scheduler
    run_scheduler()
    
    print(f"🚀 Vise-AI Modern interface: http://localhost:8000")
    # Use debug=False on Windows with Socket.IO to avoid "not a socket" errors
    socketio.run(app, host='0.0.0.0', port=8000, debug=False, allow_unsafe_werkzeug=True)
